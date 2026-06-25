from __future__ import annotations

import json
import os
from io import BytesIO
from datetime import date, datetime
from functools import wraps
from typing import Dict, List, Tuple
from types import SimpleNamespace
from flask import jsonify

from flask import (
    Flask,
    flash,
    g,
    has_request_context,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from flask_login import (
    LoginManager,
    UserMixin,
    current_user,
    login_required,
    login_user,
    logout_user,
)
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename
from openpyxl import Workbook, load_workbook

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DB_PATH = os.path.join(BASE_DIR, "instance", "shooting.db")

# ให้ SQLite เขียนได้หลังแตกไฟล์ ZIP บน Mac/Windows
def ensure_sqlite_writable():
    if os.environ.get("DATABASE_URL"):
        return
    try:
        os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
        os.chmod(os.path.dirname(DB_PATH), 0o755)
        if os.path.exists(DB_PATH):
            os.chmod(DB_PATH, 0o666)
    except Exception:
        # ไม่ให้ระบบล่มเพราะ chmod บนบาง host ไม่รองรับ
        pass

ensure_sqlite_writable()

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret-change-me")

database_url = os.environ.get("DATABASE_URL")
if database_url:
    if database_url.startswith("postgres://"):
        database_url = database_url.replace("postgres://", "postgresql://", 1)
    app.config["SQLALCHEMY_DATABASE_URI"] = database_url
else:
    app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{DB_PATH}"

app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = "login"

LANG_LABELS = {
    "th": {
        "name": "ไทย",
        "shooting_title": "ประเภทสุดยอดความแม่นยำ (SHOOTING)",
    },
    "en": {
        "name": "English",
        "shooting_title": "Precision Shooting",
    },
    "fr": {
        "name": "Français",
        "shooting_title": "Tir de précision",
    },
    "zh": {
        "name": "中文",
        "shooting_title": "精准射击",
    },
}
SUPPORTED_LANGS = tuple(LANG_LABELS.keys())


def current_language() -> str:
    lang = session.get("lang", "th")
    return lang if lang in SUPPORTED_LANGS else "th"


@app.context_processor
def inject_language_options():
    lang = current_language()
    return {
        "current_lang": lang,
        "lang_labels": LANG_LABELS,
        "html_lang": "zh-CN" if lang == "zh" else lang,
    }

ROUND_LABELS = {
    1: "รอบที่ 1",
    2: "รอบที่ 2",
}

SCORECARD_ROUND_LABELS_8 = [
    "รอบที่ 1",
    "รอบที่ 2",
    "รอบ 8 คน",
    "รอบรองชนะเลิศ",
    "รอบชิงชนะเลิศ",
]

SCORECARD_ROUND_LABELS_16 = [
    "รอบที่ 1",
    "รอบที่ 2",
    "รอบ 16 คน",
    "รอบ 8 คน",
    "รอบรองชนะเลิศ",
    "รอบชิงชนะเลิศ",
]

ALL_SCORECARD_ROUNDS = [1, 2, 3, 4, 5, 6]

STATIONS = [1, 2, 3, 4, 5]
DISTANCES = [6, 7, 8, 9]
MAX_RED_CARDS = 2


class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), nullable=False, default="user")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def set_password(self, password: str) -> None:
        self.password_hash = generate_password_hash(password)

    def check_password(self, password: str) -> bool:
        return check_password_hash(self.password_hash, password)


class Event(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), nullable=False)
    event_group = db.Column(db.String(50), nullable=False)
    category = db.Column(db.String(50), nullable=False)
    competition_date = db.Column(db.Date, nullable=False)
    location = db.Column(db.String(255), nullable=False)
    lane_count = db.Column(db.Integer, nullable=False, default=1)
    direct_qualifiers = db.Column(db.Integer, nullable=False, default=0)
    has_round_two = db.Column(db.Boolean, default=False)
    round_two_cutoff_rank = db.Column(db.Integer, nullable=True)
    next_round_label = db.Column(db.String(50), nullable=False, default="รอบ 8 คน")
    round_two_advancers = db.Column(db.Integer, nullable=False, default=4)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    created_by = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=True)

    athletes = db.relationship("Athlete", backref="event", cascade="all, delete-orphan", lazy=True)


class Athlete(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    event_id = db.Column(db.Integer, db.ForeignKey("event.id"), nullable=False)
    bib_no = db.Column(db.String(20), nullable=False)
    name = db.Column(db.String(255), nullable=False)
    affiliation = db.Column(db.String(255), nullable=False)
    start_order = db.Column(db.Integer, nullable=False)
    lane_no = db.Column(db.Integer, nullable=False)
    lane_order = db.Column(db.Integer, nullable=False)
    status = db.Column(db.String(20), nullable=False, default="waiting")
    red_card_count = db.Column(db.Integer, nullable=False, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    entries = db.relationship("ScoreEntry", backref="athlete", cascade="all, delete-orphan", lazy=True)
    signatures = db.relationship("ScoreSignature", backref="athlete", cascade="all, delete-orphan", lazy=True)
    tiebreaks = db.relationship("TieBreakEntry", backref="athlete", cascade="all, delete-orphan", lazy=True)


class ScoreEntry(db.Model):
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    athlete_id = db.Column(db.Integer, db.ForeignKey("athlete.id"), nullable=False)
    round_no = db.Column(db.Integer, nullable=False, default=1)
    station_no = db.Column(db.Integer, nullable=False)
    distance_m = db.Column(db.Integer, nullable=False)
    score = db.Column(db.Integer, nullable=False, default=0)
    is_red_card = db.Column(db.Boolean, nullable=False, default=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class ScoreSignature(db.Model):
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    athlete_id = db.Column(db.Integer, db.ForeignKey("athlete.id"), nullable=False)
    round_no = db.Column(db.Integer, nullable=False)
    recorder_name = db.Column(db.String(255), nullable=True)
    referee_name = db.Column(db.String(255), nullable=True)
    athlete_name = db.Column(db.String(255), nullable=True)
    recorder_signature = db.Column(db.Text, nullable=True)
    referee_signature = db.Column(db.Text, nullable=True)
    athlete_signature = db.Column(db.Text, nullable=True)
    bypass_signed = db.Column(db.Boolean, default=False)
    started_at = db.Column(db.DateTime, nullable=True)
    finished_at = db.Column(db.DateTime, nullable=True)


class TieBreakEntry(db.Model):
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    athlete_id = db.Column(db.Integer, db.ForeignKey("athlete.id"), nullable=False)
    round_no = db.Column(db.Integer, nullable=False)
    station_no = db.Column(db.Integer, nullable=False)
    score = db.Column(db.Integer, nullable=False, default=0)


class BracketMatch(db.Model):
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    event_id = db.Column(db.Integer, db.ForeignKey("event.id"), nullable=False)
    round_name = db.Column(db.String(20), nullable=False)
    match_no = db.Column(db.Integer, nullable=False)
    athlete_a_id = db.Column(db.Integer, nullable=True)
    athlete_b_id = db.Column(db.Integer, nullable=True)
    winner_id = db.Column(db.Integer, nullable=True)


class ResultsApprovedSetting(db.Model):
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    event_id = db.Column(db.Integer, db.ForeignKey("event.id"), unique=True, nullable=False)
    competition_title = db.Column(db.String(255), nullable=True)
    host_line = db.Column(db.String(255), nullable=True)
    date_line = db.Column(db.String(255), nullable=True)
    location_line = db.Column(db.String(255), nullable=True)
    country_label = db.Column(db.String(80), nullable=False, default="COUNTRY")
    president_title = db.Column(db.String(255), nullable=True)
    president_name = db.Column(db.String(255), nullable=True)
    technical_title = db.Column(db.String(255), nullable=True)
    technical_name = db.Column(db.String(255), nullable=True)
    umpires_text = db.Column(db.Text, nullable=True)
    approved_text = db.Column(db.String(255), nullable=False, default="……………………………APPROVED")
    show_official_pages = db.Column(db.Boolean, nullable=False, default=True)
    cover_main_logo_path = db.Column(db.String(255), nullable=True)
    cover_bottom_logo_1_path = db.Column(db.String(255), nullable=True)
    cover_bottom_logo_2_path = db.Column(db.String(255), nullable=True)
    cover_bottom_logo_3_path = db.Column(db.String(255), nullable=True)
    header_logo_1_path = db.Column(db.String(255), nullable=True)
    header_logo_2_path = db.Column(db.String(255), nullable=True)
    header_logo_3_path = db.Column(db.String(255), nullable=True)
    header_logo_4_path = db.Column(db.String(255), nullable=True)
    side_logo_path = db.Column(db.String(255), nullable=True)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


@login_manager.user_loader
def load_user(user_id: str):
    return User.query.get(int(user_id))


def role_required(*roles: str):
    def decorator(func_):
        @wraps(func_)
        def wrapper(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for("login"))
            if current_user.role == "superadmin" or current_user.role in roles:
                return func_(*args, **kwargs)
            flash("คุณไม่มีสิทธิ์ทำรายการนี้", "warning")
            return redirect(url_for("index"))
        return wrapper
    return decorator


def next_manual_id(model):
    """Fallback สำหรับ PostgreSQL table เก่าที่ id ไม่มี sequence/default"""
    try:
        max_id = db.session.query(db.func.max(model.id)).scalar() or 0
        return int(max_id) + 1
    except Exception:
        db.session.rollback()
        return None

#---------------คะแนนเรียลไทม์------------
def event_has_round_of_16(event) -> bool:
    if event.next_round_label == "รอบ 16 คน":
        return True
    return BracketMatch.query.filter_by(event_id=event.id, round_name="R16").first() is not None


def scorecard_round_labels(event) -> list[str]:
    return SCORECARD_ROUND_LABELS_16 if event_has_round_of_16(event) else SCORECARD_ROUND_LABELS_8


def scorecard_round_numbers(event) -> list[int]:
    return list(range(1, len(scorecard_round_labels(event)) + 1))


def bracket_round_to_scorecard_round(round_name: str, event=None) -> int:
    if round_name == "R16":
        return 3
    if event is not None and event_has_round_of_16(event):
        return {"QF": 4, "SF": 5, "F": 6}[round_name]
    return {"QF": 3, "SF": 4, "F": 5}[round_name]

def build_bracket_row_data(event, athlete, round_name, seed_map=None):
    seed_map = seed_map or {}
    if not athlete:
        return {
            "athlete_id": None,
            "seed": "",
            "team": "",
            "name": "",
            "r1": "",
            "r1r2": "",
            "show_ref_score": False,
            "is_direct_qualifier": False,
            "stations": [0, 0, 0, 0, 0],
            "total": 0,
            "status": "waiting",
            "status_label": "รอคิว",
        }

    round_no = bracket_round_to_scorecard_round(round_name, event)

    # ใช้ group ปัจจุบันของระบบ: direct = เข้ารอบตรงจากรอบ 1
    progression_groups = get_progression_groups(event)
    is_direct_qualifier = athlete.id in progression_groups.get("direct", set())

    r1_summary = summarize_round(athlete.id, 1)
    r2_summary = summarize_round(athlete.id, 2) if event.has_round_two else {"total": 0}

    # S1-S5 และ SCORE ของ bracket ต้องดึงจาก scorecard รอบ bracket ปัจจุบันแบบ realtime
    current_summary = summarize_round(athlete.id, round_no) if round_no else {"by_station": {}, "total": 0}

    r1_total = r1_summary.get("total", 0)
    r2_total = r2_summary.get("total", 0) or 0

    # REF = รอบ 1 + รอบ 2
    # ซ่อน REF เฉพาะผู้เข้ารอบตรงจากรอบ 1 เท่านั้น
    ref_total = r1_total + r2_total
    show_ref_score = not is_direct_qualifier

    stations = []
    by_station = current_summary.get("by_station", {})
    for station in [1, 2, 3, 4, 5]:
        station_data = by_station.get(station, {"total": 0})
        stations.append(station_data.get("total", 0))

    round_status = athlete_round_status(athlete, round_no)

    return {
        "athlete_id": athlete.id,
        "seed": seed_map.get(athlete.id, ""),
        "team": athlete.affiliation or "",
        "name": athlete.name or "",
        "r1": r1_total,
        "r1r2": ref_total if show_ref_score else "",
        "show_ref_score": show_ref_score,
        "is_direct_qualifier": is_direct_qualifier,
        "stations": stations,
        "total": current_summary.get("total", 0),
        "status": round_status,
        "status_label": {"waiting": "รอคิว", "active": "กำลังตี", "finished": "ตีเสร็จแล้ว"}.get(round_status, "รอคิว"),
    }

def ensure_schema() -> None:
    os.makedirs(os.path.join(BASE_DIR, "instance"), exist_ok=True)

    # PostgreSQL บน Railway: ตารางที่เคยย้ายจาก SQLite บางตัวมี id NOT NULL
    # แต่ไม่มี DEFAULT nextval(sequence) ทำให้ INSERT แล้ว id = null
    # แก้แบบถาวรด้วยการสร้าง sequence และผูก default ให้ทุกตารางหลัก
    if db.engine.dialect.name == "postgresql":
        tables = [
            "user",
            "event",
            "athlete",
            "score_entry",
            "score_signature",
            "tie_break_entry",
            "bracket_match",
            "results_approved_setting",
        ]
        with db.engine.begin() as conn:
            for table in tables:
                seq = f"{table}_id_seq"
                conn.exec_driver_sql(f'CREATE SEQUENCE IF NOT EXISTS "{seq}"')
                conn.exec_driver_sql(
                    f"""SELECT setval(
                        '"{seq}"',
                        COALESCE((SELECT MAX(id) FROM "{table}"), 0) + 1,
                        false
                    )"""
                )
                conn.exec_driver_sql(
                    f'ALTER TABLE "{table}" ALTER COLUMN id SET DEFAULT nextval(\'"{seq}"\')'
                )
                conn.exec_driver_sql(
                    f'ALTER SEQUENCE "{seq}" OWNED BY "{table}".id'
                )
            # เพิ่มคอลัมน์สำหรับอัปโหลดโลโก้ Results Approved ในฐานข้อมูลเดิม
            ra_logo_columns = {
                "cover_main_logo_path": "VARCHAR(255)",
                "cover_bottom_logo_1_path": "VARCHAR(255)",
                "cover_bottom_logo_2_path": "VARCHAR(255)",
                "cover_bottom_logo_3_path": "VARCHAR(255)",
                "header_logo_1_path": "VARCHAR(255)",
                "header_logo_2_path": "VARCHAR(255)",
                "header_logo_3_path": "VARCHAR(255)",
                "header_logo_4_path": "VARCHAR(255)",
                "side_logo_path": "VARCHAR(255)",
            }
            for col, col_type in ra_logo_columns.items():
                conn.exec_driver_sql(f'ALTER TABLE "results_approved_setting" ADD COLUMN IF NOT EXISTS {col} {col_type}')
        return

    # SQLite migration เดิม
    if db.engine.dialect.name != "sqlite":
        return

    with db.engine.begin() as conn:
        columns = {row[1] for row in conn.exec_driver_sql("PRAGMA table_info(score_signature)").fetchall()}
        if "recorder_signature" not in columns:
            conn.exec_driver_sql("ALTER TABLE score_signature ADD COLUMN recorder_signature TEXT")
        if "referee_signature" not in columns:
            conn.exec_driver_sql("ALTER TABLE score_signature ADD COLUMN referee_signature TEXT")
        if "athlete_signature" not in columns:
            conn.exec_driver_sql("ALTER TABLE score_signature ADD COLUMN athlete_signature TEXT")
        if "bypass_signed" not in columns:
            conn.exec_driver_sql("ALTER TABLE score_signature ADD COLUMN bypass_signed BOOLEAN DEFAULT 0")
        if "started_at" not in columns:
            conn.exec_driver_sql("ALTER TABLE score_signature ADD COLUMN started_at DATETIME")
        if "finished_at" not in columns:
            conn.exec_driver_sql("ALTER TABLE score_signature ADD COLUMN finished_at DATETIME")

        event_columns = {row[1] for row in conn.exec_driver_sql("PRAGMA table_info(event)").fetchall()}
        if "has_round_two" not in event_columns:
            conn.exec_driver_sql("ALTER TABLE event ADD COLUMN has_round_two BOOLEAN DEFAULT 1")
        if "bracket_qualifiers" not in event_columns:
            conn.exec_driver_sql("ALTER TABLE event ADD COLUMN bracket_qualifiers INTEGER DEFAULT 8")
        if "direct_qualifiers" not in event_columns:
            conn.exec_driver_sql("ALTER TABLE event ADD COLUMN direct_qualifiers INTEGER DEFAULT 4")
        if "category" not in event_columns:
            conn.exec_driver_sql("ALTER TABLE event ADD COLUMN category VARCHAR(20) DEFAULT 'men'")
        if "next_round_label" not in event_columns:
            conn.exec_driver_sql("ALTER TABLE event ADD COLUMN next_round_label VARCHAR(50)")

        athlete_columns = {row[1] for row in conn.exec_driver_sql("PRAGMA table_info(athlete)").fetchall()}
        if "lane_no" not in athlete_columns:
            conn.exec_driver_sql("ALTER TABLE athlete ADD COLUMN lane_no INTEGER")
        if "lane_order" not in athlete_columns:
            conn.exec_driver_sql("ALTER TABLE athlete ADD COLUMN lane_order INTEGER")
        if "start_order" not in athlete_columns:
            conn.exec_driver_sql("ALTER TABLE athlete ADD COLUMN start_order INTEGER")

        # Results Approved: คอลัมน์เก็บ path โลโก้ที่อัปโหลดเอง
        ra_columns = {row[1] for row in conn.exec_driver_sql("PRAGMA table_info(results_approved_setting)").fetchall()}
        ra_logo_columns = {
            "cover_main_logo_path": "VARCHAR(255)",
            "cover_bottom_logo_1_path": "VARCHAR(255)",
            "cover_bottom_logo_2_path": "VARCHAR(255)",
            "cover_bottom_logo_3_path": "VARCHAR(255)",
            "header_logo_1_path": "VARCHAR(255)",
            "header_logo_2_path": "VARCHAR(255)",
            "header_logo_3_path": "VARCHAR(255)",
            "header_logo_4_path": "VARCHAR(255)",
            "side_logo_path": "VARCHAR(255)",
        }
        for col, col_type in ra_logo_columns.items():
            if col not in ra_columns:
                conn.exec_driver_sql(f"ALTER TABLE results_approved_setting ADD COLUMN {col} {col_type}")


def event_theme(category: str | None) -> str:
    """คืนค่า class สีของ event ตามประเภทการแข่งขัน"""
    category = (category or "men").strip().lower()
    if category in {"women", "female", "หญิง", "lady", "ladies"}:
        return "theme-women"
    if category in {"mixed", "mix", "ผสม", "คู่ผสม"}:
        return "theme-mixed"
    if category in {"youth", "junior", "เยาวชน"}:
        return "theme-youth"
    return "theme-men"


def seed_defaults() -> None:
    if not User.query.filter_by(username="superadmin").first():
        u = User(username="superadmin", role="superadmin")
        u.set_password("yagami1225")
        db.session.add(u)
    if not User.query.filter_by(username="admin").first():
        u = User(username="admin", role="admin")
        u.set_password("admin1234")
        db.session.add(u)
    if not User.query.filter_by(username="viewer").first():
        u = User(username="viewer", role="user")
        u.set_password("viewer1234")
        db.session.add(u)
    db.session.commit()



def _request_cache() -> dict:
    """Request-local cache for read-heavy overview and bracket pages.

    The live polling cadence stays unchanged. This only prevents the same SQL
    queries from being repeated many times while building one response.
    """
    if not has_request_context():
        return {}
    cache = getattr(g, "shooting_request_cache", None)
    if cache is None:
        cache = {}
        g.shooting_request_cache = cache
    return cache


def clear_request_cache() -> None:
    if has_request_context():
        g.shooting_request_cache = {}


def preload_event_score_data(event: Event) -> None:
    """Load entries, signatures and tie-break scores for an event in 3 queries."""
    if not has_request_context():
        return
    cache = _request_cache()
    if ("preloaded_event", event.id) in cache:
        return
    athlete_ids = [a.id for a in event.athletes]
    cache[("preloaded_event", event.id)] = True
    if not athlete_ids:
        return

    entries_by_key = cache.setdefault("entries_by_athlete_round", {})
    for entry in ScoreEntry.query.filter(ScoreEntry.athlete_id.in_(athlete_ids)).all():
        entries_by_key.setdefault((entry.athlete_id, entry.round_no), []).append(entry)

    signatures_by_key = cache.setdefault("signature_by_athlete_round", {})
    for signature in ScoreSignature.query.filter(ScoreSignature.athlete_id.in_(athlete_ids)).all():
        signatures_by_key[(signature.athlete_id, signature.round_no)] = signature

    tiebreak_by_key = cache.setdefault("tiebreak_by_athlete_round", {})
    for entry in TieBreakEntry.query.filter(TieBreakEntry.athlete_id.in_(athlete_ids)).all():
        tiebreak_by_key.setdefault((entry.athlete_id, entry.round_no), []).append(entry)

def get_round_score_map(athlete_id: int, round_no: int) -> Dict[Tuple[int, int], ScoreEntry]:
    entries = ScoreEntry.query.filter_by(athlete_id=athlete_id, round_no=round_no).all()
    return {(e.station_no, e.distance_m): e for e in entries}


def get_round_signature(athlete_id: int, round_no: int) -> ScoreSignature | None:
    cache = _request_cache()
    signatures = cache.get("signature_by_athlete_round") if has_request_context() else None
    key = (athlete_id, round_no)
    if signatures is not None and key in signatures:
        return signatures[key]
    signature = ScoreSignature.query.filter_by(athlete_id=athlete_id, round_no=round_no).first()
    if has_request_context():
        cache.setdefault("signature_by_athlete_round", {})[key] = signature
    return signature


def ensure_round_entries(athlete_id: int, round_no: int) -> None:
    existing = {(e.station_no, e.distance_m) for e in ScoreEntry.query.filter_by(athlete_id=athlete_id, round_no=round_no).all()}
    changed = False
    for station_no in STATIONS:
        for distance_m in DISTANCES:
            key = (station_no, distance_m)
            if key not in existing:
                db.session.add(ScoreEntry(
                    athlete_id=athlete_id,
                    round_no=round_no,
                    station_no=station_no,
                    distance_m=distance_m,
                    score=0,
                    is_red_card=False,
                ))
                changed = True
    if changed:
        db.session.commit()
        clear_request_cache()


def ensure_signature(athlete_id: int, round_no: int) -> ScoreSignature:
    signature = get_round_signature(athlete_id, round_no)
    if signature:
        return signature
    signature = ScoreSignature(athlete_id=athlete_id, round_no=round_no)
    db.session.add(signature)
    db.session.commit()
    clear_request_cache()
    return signature


def summarize_round(athlete_id: int, round_no: int) -> dict:
    cache = _request_cache()
    summary_cache = cache.setdefault("round_summary", {}) if has_request_context() else {}
    key = (athlete_id, round_no)
    if key in summary_cache:
        return summary_cache[key]

    entries_map = cache.get("entries_by_athlete_round") if has_request_context() else None
    if entries_map is not None:
        entries = entries_map.get(key, [])
    else:
        entries = ScoreEntry.query.filter_by(athlete_id=athlete_id, round_no=round_no).all()

    total = sum(e.score for e in entries)
    count_5 = sum(1 for e in entries if e.score == 5)
    count_3 = sum(1 for e in entries if e.score == 3)
    red_cards = sum(1 for e in entries if e.is_red_card)
    by_station = {}
    for station_no in STATIONS:
        station_entries = [e for e in entries if e.station_no == station_no]
        by_station[station_no] = {
            "distances": {e.distance_m: e.score for e in station_entries},
            "total": sum(e.score for e in station_entries),
        }

    tiebreak_map = cache.get("tiebreak_by_athlete_round") if has_request_context() else None
    if tiebreak_map is not None:
        tiebreak_entries = tiebreak_map.get(key, [])
    else:
        tiebreak_entries = TieBreakEntry.query.filter_by(athlete_id=athlete_id, round_no=round_no).all()
    result = {
        "total": total,
        "count_5": count_5,
        "count_3": count_3,
        "red_cards": red_cards,
        "by_station": by_station,
        "tiebreak_total": sum(e.score for e in tiebreak_entries),
        "tiebreak_count": len(tiebreak_entries),
    }
    summary_cache[key] = result
    return result


def build_scorecard_template_data(athlete_id: int) -> dict:
    entries = ScoreEntry.query.filter_by(athlete_id=athlete_id).all()

    score_map: dict[tuple[int, int, int], dict] = {}
    station_totals: dict[tuple[int, int], int] = {}
    station_reds: dict[tuple[int, int], int] = {}
    round_totals: dict[int, int] = {}

    for round_no in ALL_SCORECARD_ROUNDS:
        round_totals[round_no] = 0
        for station_no in STATIONS:
            station_entries = [
                e for e in entries
                if e.round_no == round_no and e.station_no == station_no
            ]
            station_total = sum(e.score for e in station_entries)
            station_red = sum(1 for e in station_entries if e.is_red_card)

            station_totals[(round_no, station_no)] = station_total
            station_reds[(round_no, station_no)] = station_red
            round_totals[round_no] += station_total

            for distance_m in DISTANCES:
                entry = next(
                    (e for e in station_entries if e.distance_m == distance_m),
                    None
                )
                score_map[(round_no, station_no, distance_m)] = {
                    "score": entry.score if entry else "",
                    "is_red_card": entry.is_red_card if entry else False,
                }

    round_ranks = {1: "", 2: "", 3: "", 4: "", 5: ""}
    return {
        "score_map": score_map,
        "station_totals": station_totals,
        "station_reds": station_reds,
        "round_totals": round_totals,
        "round_ranks": round_ranks,
    }


def athlete_round_status(athlete: Athlete, round_no: int) -> str:
    signature = get_round_signature(athlete.id, round_no)
    if signature and signature.finished_at:
        return "finished"
    if signature and signature.started_at:
        return "active"
    return "waiting"


def bracket_match_status(event: Event, match: BracketMatch) -> dict:
    """สถานะของคู่ใน bracket จากลายเซ็น/เวลาเริ่มของนักกีฬาทั้ง 2 ฝั่ง"""
    round_no = bracket_round_to_scorecard_round(match.round_name, event)
    athlete_ids = [aid for aid in [match.athlete_a_id, match.athlete_b_id] if aid]
    if not athlete_ids or len(athlete_ids) < 2:
        return {"key": "waiting", "label": "รอคู่แข่งขัน"}
    if match.winner_id:
        return {"key": "finished", "label": "แข่งเสร็จ"}
    statuses = []
    for aid in athlete_ids:
        sig = get_round_signature(aid, round_no)
        if sig and sig.finished_at:
            statuses.append("finished")
        elif sig and sig.started_at:
            statuses.append("active")
        else:
            statuses.append("waiting")
    if any(status == "active" for status in statuses):
        return {"key": "active", "label": "กำลังแข่งขัน"}
    if all(status == "finished" for status in statuses):
        return {"key": "ready_result", "label": "รอบันทึกผู้ชนะ"}
    if any(status == "finished" for status in statuses):
        return {"key": "active", "label": "กำลังแข่งขัน"}
    return {"key": "waiting", "label": "รอแข่งขัน"}


def ranking_key(item: dict):
    """
    สูตรจัดอันดับ Shooting ที่ใช้ในระบบ
    1) คะแนนรวม / TOTAL มากกว่า
    2) จำนวนคะแนน 5 มากกว่า
    3) จำนวนคะแนน 3 มากกว่า
    4) คะแนน Shoot-off 7 เมตรทุกสถานี (เมื่อมีการบันทึก)

    หมายเหตุ: ตัดคะแนนสถานี 5 และสถานี 4 ออกจากสูตรแล้ว
    เพื่อไม่ให้สถานีใดมีน้ำหนักมากกว่าสถานีอื่น
    """
    return (
        item["total"],
        item["count_5"],
        item["count_3"],
        item["tiebreak_total"],
    )


def apply_rank_by_tiebreak(rows: list[dict]) -> None:
    """ใส่อันดับแบบต่อเนื่อง 1..คนสุดท้าย หลัง sort ด้วยกติกาจริงแล้ว

    กติกาใหม่ของระบบ: หน้า Overview ห้ามแสดง Class/Rank ซ้ำเอง
    ถ้า TOTAL -> จำนวน 5 -> จำนวน 3 ยังเท่ากันในช่วงที่มีผล
    จะให้ระบบขึ้น Shoot-off เพื่อแยกอันดับแทน
    """
    for idx, row in enumerate(rows, start=1):
        row["rank"] = idx
        row["ordinal_rank"] = idx
        row["display_rank"] = idx
        row["view_order"] = idx




def bracket_size(event: Event) -> int:
    """จำนวนคนรอบ Knockout จาก label เช่น รอบ 16/8/4"""
    label = (event.next_round_label or "")
    if "16" in label:
        return 16
    if "4" in label or "รอง" in label:
        return 4
    return 8


def direct_quota(event: Event) -> int:
    """จำนวนคนที่เข้ารอบตรงจากรอบแรก

    ต้องยึดค่าที่ผู้ใช้กรอกในอีเวนต์ก่อน เช่น
    - รอบ 16 กรอกผ่านรอบแรก 8 = ต้องเป็น 8 จริง
    - รอบ 8 กรอกผ่านรอบแรก 4 = ต้องเป็น 4 จริง

    ถ้าอีเวนต์เก่าไม่ได้กรอก direct_qualifiers จึงค่อย fallback เป็นครึ่งหนึ่งของ bracket
    """
    size = bracket_size(event)
    configured = int(event.direct_qualifiers or 0)
    if configured > 0:
        return min(configured, size)
    if event.has_round_two:
        return max(size // 2, 0)
    return size


def round2_advancer_quota(event: Event) -> int:
    """จำนวนคนที่ผ่านจากรอบ 2 เข้า Knockout

    ต้องยึดค่าที่ผู้ใช้กรอกจริง เช่น รอบ 16 ถ้ากรอกให้ผ่านจากรอบ 2 = 8
    ก็ต้องใช้ 8 ไม่ใช่คำนวณเหลือเองจนเพี้ยน
    """
    if not event.has_round_two:
        return 0
    size = bracket_size(event)
    configured = int(event.round_two_advancers or 0)
    remaining = max(size - direct_quota(event), 0)
    if configured > 0:
        return min(configured, size)
    return remaining


def round1_total_cut_ids(rows: list[dict], cutoff_count: int) -> set[int]:
    """สิทธิ์ตีรอบ 2 ใช้กติกาตามเอกสาร: ถ้าคะแนนรวมเท่ากับคนลำดับสุดท้าย ได้ตีรอบ 2 ทั้งหมด
    ไม่ใช้ 5/3 มาตัดสิทธิ์ตีรอบสอง เพราะรอบสองคือการคัดต่อ ไม่ใช่ตัดเข้า Knockout
    """
    if not cutoff_count or cutoff_count <= 0:
        return set()
    if not rows:
        return set()
    if len(rows) <= cutoff_count:
        return {row["athlete"].id for row in rows}
    cutoff_total = rows[cutoff_count - 1].get("total", 0)
    return {row["athlete"].id for row in rows if row.get("total", 0) >= cutoff_total}


def apply_round1_round2_cutoff_display_rank(event: Event, rows: list[dict]) -> None:
    """แสดงลำดับซ้ำได้เฉพาะกลุ่มสุดท้ายที่ได้สิทธิ์ตีรอบ 2

    กติกา: ภาพรวมรอบ 1 ต้องเรียง 1..คนสุดท้าย ไม่ให้ซ้ำเอง
    ยกเว้นถ้ามีรอบ 2 และคะแนนรวมเท่ากับคนสุดท้ายของสิทธิ์รอบ 2
    กลุ่มนั้นให้ได้สิทธิ์ตีรอบ 2 ทั้งหมด และแสดง Class เป็นเลขเส้นตัดเดียวกันได้
    ส่วนลำดับจริงยังเก็บไว้ใน ordinal_rank สำหรับเรียงแถว/คำนวณต่อ
    """
    if not rows:
        return
    for idx, row in enumerate(rows, start=1):
        row["rank"] = idx
        row["ordinal_rank"] = idx
        row["display_rank"] = idx
        row["view_order"] = idx

    if not event.has_round_two or not event.round_two_cutoff_rank:
        return

    cutoff = int(event.round_two_cutoff_rank or 0)
    direct = direct_quota(event)
    if cutoff <= direct or cutoff <= 0 or len(rows) < cutoff:
        return

    cutoff_total = rows[cutoff - 1].get("total", 0)
    final_total_indexes = [
        idx for idx, row in enumerate(rows)
        if idx >= direct and row.get("total", 0) == cutoff_total
    ]
    if len(final_total_indexes) < 2 or (cutoff - 1) not in final_total_indexes:
        return

    for idx in final_total_indexes:
        rows[idx]["rank"] = cutoff
        rows[idx]["display_rank"] = cutoff
        rows[idx]["view_order"] = rows[idx].get("ordinal_rank", idx + 1)


def apply_sequential_rank(rows: list[dict], start: int = 1) -> None:
    """ใช้กับ Overview รอบ 2/Seed: ต้องไม่มีลำดับซ้ำ"""
    for idx, row in enumerate(rows, start=start):
        row["rank"] = idx
        row["ordinal_rank"] = idx
        row["display_rank"] = idx
        row["view_order"] = idx

def build_round_ranking(event: Event, round_no: int) -> List[dict]:
    cache = _request_cache()
    cache_key = ("round_ranking", event.id, round_no)
    if has_request_context() and cache_key in cache:
        return cache[cache_key]

    preload_event_score_data(event)
    athletes = sorted(event.athletes, key=lambda a: a.start_order)
    rows = []
    round2_display_map = {}
    if round_no == 2:
        round1_rows = build_round_ranking(event, 1)
        direct_ids_for_r2 = exact_cut_ids(round1_rows, direct_quota(event))
        direct_pending_ids_for_r2 = unresolved_tie_ids(round1_rows, direct_quota(event))
        round2_source_rows = [
            row for row in round1_rows
            if row["athlete"].id not in direct_ids_for_r2
            and row["athlete"].id not in direct_pending_ids_for_r2
            and is_round_two_candidate(event, row["athlete"])
        ]
        round2_source_rows.sort(key=lambda row: (
            row["total"], row["count_5"], row["count_3"], row["tiebreak_total"], row["athlete"].start_order,
        ))
        lane_count = max(event.lane_count or 1, 1)
        for idx, source_row in enumerate(round2_source_rows, start=1):
            round2_display_map[source_row["athlete"].id] = {
                "display_order": idx,
                "display_lane_no": ((idx - 1) % lane_count) + 1,
                "display_lane_order": ((idx - 1) // lane_count) + 1,
            }

    for athlete in athletes:
        if round_no == 2 and not get_round_signature(athlete.id, 2):
            continue
        summary = summarize_round(athlete.id, round_no)
        row = {
            "athlete": athlete,
            "round_no": round_no,
            "total": summary["total"],
            "count_5": summary["count_5"],
            "count_3": summary["count_3"],
            "tiebreak_total": summary["tiebreak_total"],
            "tiebreak_count": summary.get("tiebreak_count", 0),
            "status": athlete_round_status(athlete, round_no),
            "by_station": summary["by_station"],
            "red_cards": summary["red_cards"],
            "display_order": athlete.start_order,
            "display_lane_no": athlete.lane_no,
            "display_lane_order": athlete.lane_order,
        }
        if round_no == 2 and athlete.id in round2_display_map:
            row.update(round2_display_map[athlete.id])
        rows.append(row)
    rows.sort(key=ranking_key, reverse=True)
    apply_rank_by_tiebreak(rows)
    if round_no == 1:
        apply_round1_round2_cutoff_display_rank(event, rows)
    if has_request_context():
        cache[cache_key] = rows
    return rows


def round_two_candidate_ids(event: Event) -> set[int]:
    cache = _request_cache()
    cache_key = ("round_two_candidate_ids", event.id)
    if has_request_context() and cache_key in cache:
        return cache[cache_key]

    ids: set[int] = set()
    if event.has_round_two and event.round_two_cutoff_rank:
        cutoff = int(event.round_two_cutoff_rank or 0)
        direct = direct_quota(event)
        round1_rows = build_round_ranking(event, 1)

        # สิทธิ์ตีรอบ 2 ตามกติกา: คนถัดจากเข้ารอบตรงจนถึงลำดับที่ตั้งไว้
        # ถ้าคนลำดับสุดท้ายของสิทธิ์รอบ 2 คะแนนรวมเท่ากัน ให้ได้ตีรอบ 2 ทั้งหมด
        # แต่ถ้ามี Shoot-off ที่เส้นเข้ารอบตรงอยู่ ให้รอผลก่อน ไม่ดันกลุ่มนั้นไปเป็นรอบ 2
        direct_ids = exact_cut_ids(round1_rows, direct)
        direct_shoot_ids = unresolved_tie_ids(round1_rows, direct)
        cutoff_ids = round1_total_cut_ids(round1_rows, cutoff)
        ids = cutoff_ids - direct_ids - direct_shoot_ids

    if has_request_context():
        cache[cache_key] = ids
    return ids


def is_round_two_candidate(event: Event, athlete: Athlete) -> bool:
    return athlete.id in round_two_candidate_ids(event)


def sync_round_two_candidates(event: Event) -> None:
    if not event.has_round_two:
        return

    candidate_ids = round_two_candidate_ids(event)
    athlete_ids = [athlete.id for athlete in event.athletes]
    if not athlete_ids:
        return

    signatures = {
        sig.athlete_id: sig
        for sig in ScoreSignature.query.filter(
            ScoreSignature.athlete_id.in_(athlete_ids),
            ScoreSignature.round_no == 2,
        ).all()
    }
    existing_entry_keys = {
        (entry.athlete_id, entry.station_no, entry.distance_m)
        for entry in ScoreEntry.query.filter(
            ScoreEntry.athlete_id.in_(athlete_ids),
            ScoreEntry.round_no == 2,
        ).all()
    }

    changed = False
    for athlete in event.athletes:
        sig2 = signatures.get(athlete.id)
        if athlete.id in candidate_ids:
            for station_no in STATIONS:
                for distance_m in DISTANCES:
                    key = (athlete.id, station_no, distance_m)
                    if key not in existing_entry_keys:
                        db.session.add(ScoreEntry(
                            athlete_id=athlete.id,
                            round_no=2,
                            station_no=station_no,
                            distance_m=distance_m,
                            score=0,
                            is_red_card=False,
                        ))
                        changed = True
            if not sig2:
                sig2 = ScoreSignature(athlete_id=athlete.id, round_no=2)
                db.session.add(sig2)
                signatures[athlete.id] = sig2
                changed = True
            if not sig2.started_at and not sig2.finished_at and athlete.status != "waiting":
                athlete.status = "waiting"
                changed = True
        elif sig2 and not sig2.started_at and not sig2.finished_at:
            ScoreEntry.query.filter_by(athlete_id=athlete.id, round_no=2).delete(synchronize_session=False)
            TieBreakEntry.query.filter_by(athlete_id=athlete.id, round_no=2).delete(synchronize_session=False)
            ScoreSignature.query.filter_by(athlete_id=athlete.id, round_no=2).delete(synchronize_session=False)
            changed = True

    if changed:
        db.session.commit()
        clear_request_cache()


def build_round_two_start_list(event: Event) -> list[dict]:
    round1_rows = build_round_ranking(event, 1)

    direct_ids = exact_cut_ids(round1_rows, direct_quota(event))
    direct_pending_ids = unresolved_tie_ids(round1_rows, direct_quota(event))

    candidates = [
        r for r in round1_rows
        if is_round_two_candidate(event, r["athlete"])
        and r["athlete"].id not in direct_ids
        and r["athlete"].id not in direct_pending_ids
    ]

    # คะแนนรอบ 1 น้อยสุด ได้ตีก่อน
    candidates.sort(key=lambda r: (
        r["total"],
        r["count_5"],
        r["count_3"],
        r["tiebreak_total"],
        r["athlete"].start_order if r["athlete"].start_order is not None else 9999
    ))

    return candidates


def sort_round_two_rows_for_start(rows: list[dict]) -> list[dict]:
    """เรียงลำดับการตีรอบ 2 ตาม display_order: คะแนนรอบ 1 น้อยสุดตีก่อน"""
    return sorted(
        rows,
        key=lambda row: (
            row.get("display_order") if row.get("display_order") is not None else 999999,
            row["athlete"].start_order if row["athlete"].start_order is not None else 999999,
            row["athlete"].id,
        )
    )

def build_combined_qualifiers(event: Event) -> List[dict]:
    cache = _request_cache()
    cache_key = ("combined_qualifiers", event.id)
    if has_request_context() and cache_key in cache:
        return cache[cache_key]
    round1_rows = build_round_ranking(event, 1)
    direct_ids = exact_cut_ids(round1_rows, direct_quota(event))
    direct_pending_ids = unresolved_tie_ids(round1_rows, direct_quota(event))
    direct_rows = []
    for row in round1_rows:
        if row["athlete"].id in direct_ids:
            direct_rows.append({
                **row,
                "round1_total": row["total"],
                "round2_total": None,
                "round1_by_station": row.get("by_station", {}),
                "round2_by_station": None,
                "combined_total": row["total"],
            })
    for idx, row in enumerate(direct_rows, start=1):
        row["seed"] = idx
    if not event.has_round_two:
        return direct_rows

    candidate_ids = round_two_candidate_ids(event) - direct_ids - direct_pending_ids
    finished_candidate_ids = {
        aid for aid in candidate_ids
        if athlete_round_status(Athlete.query.get(aid), 2) == "finished"
    }
    all_round2_candidates_finished = bool(candidate_ids) and candidate_ids == finished_candidate_ids

    # รวมผลรอบ 2 เพื่อใช้คัดเข้า Bracket เฉพาะคนที่ตีจบแล้วเท่านั้น
    # กันไม่ให้คนที่ระบบสร้าง Signature ไว้แต่ยังไม่ได้ตี ถูกจัดผ่าน/ตกรอบก่อนเวลา
    round2_rows = [
        row for row in build_round_ranking(event, 2)
        if row["athlete"].id in finished_candidate_ids
    ]
    combined_rows = []
    for row in round2_rows:
        round1_row = next((r for r in round1_rows if r["athlete"].id == row["athlete"].id), None)
        round1_total = round1_row["total"] if round1_row else 0
        combined_rows.append({
            **row,
            "total": round1_total + row["total"],
            "combined_total": round1_total + row["total"],
            "count_5": (round1_row["count_5"] if round1_row else 0) + row["count_5"],
            "count_3": (round1_row["count_3"] if round1_row else 0) + row["count_3"],
            # เมื่อคัดจากรอบ 2 ใช้ SUM(R1+R2) -> 5รวม -> 3รวม -> Shoot-off รอบ 2
            # ไม่บวก Shoot-off รอบ 1 เข้ามา เพราะรอบ 1 เป็นคนละเหตุผล/คนละเที่ยว
            "tiebreak_total": row["tiebreak_total"],
            "tiebreak_count": row.get("tiebreak_count", 0),
            "round1_tiebreak_total": (round1_row.get("tiebreak_total", 0) if round1_row else 0),
            "round1_tiebreak_count": (round1_row.get("tiebreak_count", 0) if round1_row else 0),
            "round2_tiebreak_total": row["tiebreak_total"],
            "round2_tiebreak_count": row.get("tiebreak_count", 0),
            "round1_total": round1_total,
            "round2_total": row["total"],
            "round1_by_station": (round1_row.get("by_station", {}) if round1_row else {}),
            "round2_by_station": row.get("by_station", {}),
        })
    combined_rows.sort(key=ranking_key, reverse=True)
    apply_sequential_rank(combined_rows, start=1)

    # คัดผู้ผ่านจากรอบ 2 ต้องเอาตามจำนวนที่ตั้งไว้จริง ๆ
    # ถ้าคะแนนเท่ากันที่เส้นตัด ให้ตัดสินตาม TOTAL -> จำนวน 5 -> จำนวน 3 -> Shoot-off
    # ห้ามขยาย bracket เองเพราะจะทำให้ตั้ง 8 คน แต่หลุดเป็น 9/16 คน
    advancer_limit = max(round2_advancer_quota(event), 0)

    def base_tie_key(row):
        return (row["combined_total"], row["count_5"], row["count_3"])

    if all_round2_candidates_finished:
        shoot_off_ids = unresolved_tie_ids(combined_rows, advancer_limit)
        passed_ids = exact_cut_ids(combined_rows, advancer_limit)
    else:
        shoot_off_ids = set()
        passed_ids = set()

    seed_no = len(direct_rows) + 1
    for row in combined_rows:
        aid = row["athlete"].id
        row["shoot_off_required"] = aid in shoot_off_ids
        if row["shoot_off_required"]:
            row["seed"] = None
            row["passed_cut"] = False
        elif aid in passed_ids:
            row["seed"] = seed_no
            row["passed_cut"] = True
            seed_no += 1
        else:
            row["seed"] = None
            row["passed_cut"] = False
    result = direct_rows + combined_rows
    if has_request_context():
        cache[cache_key] = result
    return result




def base_shootoff_key(row: dict) -> tuple:
    """เกณฑ์ก่อนรอบพิเศษ: TOTAL -> จำนวน 5 -> จำนวน 3

    ถ้า 3 ตัวนี้ยังเท่ากัน แปลว่า "ยังจัดลำดับจริงไม่ได้"
    ต้องยิง Shoot-off ยกเว้นกรณีเส้นสุดท้ายของสิทธิ์ไปตีรอบ 2 ซึ่งกติกาให้ไปตีได้ทั้งหมด
    """
    return (
        row.get("combined_total", row.get("total", 0)),
        row.get("count_5", 0),
        row.get("count_3", 0),
    )


def tiebreak_done_key(row: dict) -> tuple:
    """ใช้ดูว่ารอบพิเศษตัดสินได้หรือยัง

    ต้องบันทึกจำนวนเที่ยวเท่ากันทุกคนก่อน แล้วจึงเอาคะแนนรอบพิเศษมาแยกอันดับ
    """
    return (row.get("tiebreak_count", 0), row.get("tiebreak_total", 0))


def cutoff_shootoff_ids(rows: list[dict], cutoff_count: int) -> set[int]:
    """คืน athlete_id ที่ต้อง Shoot-off เฉพาะเส้นตัดจริง

    เงื่อนไข:
    1) ตัดตามจำนวนที่ตั้งไว้จริง เช่น 8 คนคือ 8 คน
    2) ถ้าอันดับสุดท้ายกับคนถัดไปไม่เท่ากันตาม TOTAL -> 5 -> 3 = ไม่ต้อง Shoot-off
    3) ถ้าเท่ากัน ให้ทั้งกลุ่มที่มี key เดียวกันต้อง Shoot-off พร้อมกัน
    4) ถ้าบันทึก Shoot-off ครบทุกคนแล้วและคะแนนเที่ยวพิเศษต่างกัน = ตัดสินได้ ไม่ขึ้น Shoot-off
    5) ถ้าบันทึกยังไม่ครบ หรือบันทึกครบแล้วยังเท่ากัน = ยังขึ้น Shoot-off เพื่อให้ตีต่อ
    """
    if not cutoff_count or cutoff_count <= 0 or len(rows) <= cutoff_count:
        return set()

    cut_row = rows[cutoff_count - 1]
    next_row = rows[cutoff_count]
    base_key = base_shootoff_key(cut_row)
    if base_key != base_shootoff_key(next_row):
        return set()

    group = [row for row in rows if base_shootoff_key(row) == base_key]
    if len(group) < 2:
        return set()

    # กลุ่มต้องคร่อมเส้นตัดเท่านั้น ไม่ใช่เสมอกันภายในกลุ่มที่ผ่านหมด/ตกรอบหมด
    group_ids = {row["athlete"].id for row in group}
    above_ids = {row["athlete"].id for row in rows[:cutoff_count]}
    below_ids = {row["athlete"].id for row in rows[cutoff_count:]}
    if not (group_ids & above_ids and group_ids & below_ids):
        return set()

    counts = [row.get("tiebreak_count", 0) for row in group]
    totals = [row.get("tiebreak_total", 0) for row in group]

    # ต้องตีพร้อมกันทุกคนในกลุ่ม ถ้ามีคนใดยังไม่ได้บันทึก หรือจำนวนเที่ยวไม่เท่ากัน ให้ยังขึ้นทั้งกลุ่ม
    if min(counts) == 0 or len(set(counts)) > 1:
        return group_ids

    # บันทึกครบเท่ากันแล้ว ถ้าคะแนน Shoot-off ยังเท่ากัน ให้ตีต่อ
    if len(set(totals)) == 1:
        return group_ids

    # คะแนน Shoot-off ต่างกันแล้ว ตัดสินได้
    return set()


def unresolved_tie_ids(rows: list[dict], scope_count: int | None = None) -> set[int]:
    """หากเสมอหลัง TOTAL -> 5 -> 3 แล้วต้องเรียงลำดับจริง ให้ส่งไป Shoot-off

    scope_count ระบุจำนวนอันดับที่มีผล เช่น direct quota หรือจำนวนที่ผ่านจากรอบ 2
    ถ้ากลุ่มเสมอแตะอยู่ใน scope จะถือว่ายังตัดสินอันดับไม่ได้
    ข้อยกเว้นสิทธิ์ตีรอบ 2 รอบแรกยังใช้ round1_total_cut_ids แยกต่างหาก
    """
    if not rows:
        return set()
    scoped_ids = None
    if scope_count is not None:
        if scope_count <= 0:
            return set()
        scoped_ids = {row["athlete"].id for row in rows[:min(scope_count, len(rows))]}

    groups: dict[tuple, list[dict]] = {}
    for row in rows:
        groups.setdefault(base_shootoff_key(row), []).append(row)

    result: set[int] = set()
    for group in groups.values():
        if len(group) < 2:
            continue
        group_ids = {row["athlete"].id for row in group}
        if scoped_ids is not None and not (group_ids & scoped_ids):
            continue

        counts = [row.get("tiebreak_count", 0) for row in group]
        totals = [row.get("tiebreak_total", 0) for row in group]

        # ยังไม่ได้ตีครบทุกคน หรือจำนวนเที่ยวไม่เท่ากัน = ต้อง Shoot-off / ตีต่อ
        if min(counts) == 0 or len(set(counts)) > 1:
            result.update(group_ids)
            continue

        # ตีครบเท่ากันแล้วแต่คะแนนพิเศษยังเท่ากัน = ต้องตีต่อ
        if len(set(totals)) == 1:
            result.update(group_ids)

    return result




def overview_unresolved_shootoff_ids(rows: list[dict], scope_count: int | None = None, round_no: int | None = None) -> set[int]:
    """หาแถวที่ต้องยิง Shoot-off ในหน้า Overview

    กติกาที่ใช้:
    - เรียงด้วย TOTAL -> จำนวน 5 -> จำนวน 3 ทันที
    - ถ้า 3 ตัวนี้ยังเท่ากัน แปลว่ายังจัดลำดับจริงไม่ได้
    - ถ้ากลุ่มนั้นอยู่ในช่วงที่ต้องใช้จัดอันดับ/เข้า seed ให้ขึ้นปุ่ม Shoot-off
    - ต้องกดจบการตีของทุกคนในกลุ่มก่อน จึงขึ้น Shoot-off
    - ถ้าบันทึก Shoot-off ไม่ครบทุกคน หรือจำนวนเที่ยวไม่เท่ากัน ยังถือว่ายังตัดสินไม่ได้
    """
    if not rows:
        return set()

    scoped_ids = None
    if scope_count is not None:
        if scope_count <= 0:
            return set()
        # กลุ่มที่ชนหรืออยู่ในช่วง scope ต้องถูกตรวจทั้งกลุ่ม ไม่ใช่แค่คนใน slice
        scoped_ids = {row["athlete"].id for row in rows[:min(scope_count, len(rows))]}

    groups: dict[tuple, list[dict]] = {}
    for row in rows:
        groups.setdefault(base_shootoff_key(row), []).append(row)

    result: set[int] = set()
    for group in groups.values():
        if len(group) < 2:
            continue
        group_ids = {row["athlete"].id for row in group}
        if scoped_ids is not None and not (group_ids & scoped_ids):
            continue
        if round_no is not None:
            # ขึ้นเมื่อคนในกลุ่มนี้ตีจบครบ ไม่ต้องรอทั้งตาราง
            if not all(athlete_round_status(row["athlete"], round_no) == "finished" for row in group):
                continue

        counts = [row.get("tiebreak_count", 0) for row in group]
        totals = [row.get("tiebreak_total", 0) for row in group]

        # ยังไม่มีผลรอบพิเศษครบทุกคน หรือจำนวนเที่ยวไม่เท่ากัน = ต้องยิง/ยิงต่อ
        if min(counts) == 0 or len(set(counts)) > 1:
            result.update(group_ids)
            continue

        # ยิงรอบพิเศษครบเท่ากันแล้ว แต่คะแนนพิเศษยังเท่ากัน = ต้องยิงต่อ
        if len(set(totals)) == 1:
            result.update(group_ids)

    return result

def exact_cut_ids(rows: list[dict], cutoff_count: int) -> set[int]:
    """เลือกผู้ผ่านแบบจำนวนตายตัว และไม่ปล่อยอันดับที่ยังเสมอเข้าไปก่อน

    ใช้กับเข้ารอบตรง/ผ่านจากรอบ 2/สร้าง bracket:
    ต้องเรียงลำดับจริงด้วย TOTAL -> 5 -> 3 -> Shoot-off
    ถ้ากลุ่มเสมอแตะตำแหน่งที่มีผล ให้รอ Shoot-off ก่อน

    ข้อยกเว้นสิทธิ์ตีรอบ 2 จากรอบแรกไม่ได้ใช้ฟังก์ชันนี้ แต่ใช้ round1_total_cut_ids
    เพื่อให้คนที่คะแนนรวมเท่าลำดับสุดท้ายของสิทธิ์รอบ 2 ได้ตีทั้งหมดตามเอกสาร
    """
    if not cutoff_count or cutoff_count <= 0:
        return set()
    if not rows:
        return set()
    if len(rows) <= cutoff_count:
        pending = unresolved_tie_ids(rows, len(rows))
        return {row["athlete"].id for row in rows if row["athlete"].id not in pending}
    pending = unresolved_tie_ids(rows, cutoff_count)
    return {row["athlete"].id for row in rows[:cutoff_count] if row["athlete"].id not in pending}

def shootoff_group_ids(rows: list[dict], athlete_id: int, round_no: int | None = None) -> list[int]:
    target = next((row for row in rows if row["athlete"].id == athlete_id), None)
    if not target:
        return [athlete_id]
    key = base_shootoff_key(target)
    group = [row for row in rows if base_shootoff_key(row) == key]
    if round_no is not None:
        # ปุ่ม Shoot-off ต้องส่งเฉพาะคนในกลุ่มที่ตีจบรอบนั้นแล้ว
        # กันรอบ 2 ไปพ่วงคนรอคิว/คนเข้ารอบตรง ทำให้บันทึกแล้วดูเหมือนไม่เข้า
        group = [row for row in group if athlete_round_status(row["athlete"], round_no) == "finished"]
    ids = [row["athlete"].id for row in group]
    return ids or [athlete_id]


def _all_rows_finished_for_round(rows: list[dict], round_no: int) -> bool:
    """กันไม่ให้ขึ้น Shoot-off ก่อนกรรมการกดบันทึก/จบการตี"""
    if not rows:
        return False
    return all(athlete_round_status(row["athlete"], round_no) == "finished" for row in rows if row.get("athlete"))


def round1_overview_unresolved_shootoff_ids(event: Event, rows: list[dict]) -> set[int]:
    """หา Shoot-off รอบ 1 ตามกติกาใหม่ของครูรัก

    - ตรวจเฉพาะช่วงที่มีผล: ตั้งแต่อันดับบนสุดถึงเส้นสิทธิ์ตีรอบ 2
      ถ้าไม่มีรอบ 2 ให้ตรวจเฉพาะเส้นเข้ารอบตรง/Knockout
    - ถ้า TOTAL -> 5 -> 3 ยังเท่ากันในช่วงนั้น ให้ Shoot-off เพื่อแยกอันดับ
    - ยกเว้นกลุ่มคะแนนรวมเท่าคนสุดท้ายของสิทธิ์ตีรอบ 2 ให้ได้ตีรอบ 2 ทั้งหมด
    - คนที่ต่ำกว่าเส้นรอบ 2 แล้ว ไม่ต้อง Shoot-off
    """
    if not rows:
        return set()

    direct = direct_quota(event)
    if event.has_round_two and event.round_two_cutoff_rank:
        scope_count = int(event.round_two_cutoff_rank or 0)
    else:
        scope_count = direct
    if scope_count <= 0:
        return set()
    scope_count = min(scope_count, len(rows))

    final_round2_total = None
    if event.has_round_two and event.round_two_cutoff_rank and len(rows) >= scope_count and scope_count > direct:
        final_round2_total = rows[scope_count - 1].get("total", 0)

    groups: dict[tuple, list[tuple[int, dict]]] = {}
    for idx, row in enumerate(rows):
        groups.setdefault(base_shootoff_key(row), []).append((idx, row))

    result: set[int] = set()
    for indexed_group in groups.values():
        if len(indexed_group) < 2:
            continue
        positions = [idx for idx, _ in indexed_group]
        if not any(idx < scope_count for idx in positions):
            # ต่ำกว่าเส้นรอบ 2/เส้นเข้ารอบแล้ว ไม่ต้อง Shoot-off
            continue

        group_rows = [row for _, row in indexed_group]

        # ข้อยกเว้นเดียว: กลุ่มคะแนนรวมเท่าคนสุดท้ายของสิทธิ์ตีรอบ 2
        # และไม่ได้คร่อมเส้นเข้ารอบตรง ให้ได้ตีรอบ 2 ทั้งหมด ไม่ต้อง Shoot-off
        if final_round2_total is not None:
            same_last_round2_total = all(row.get("total", 0) == final_round2_total for row in group_rows)
            touches_round2_last_line = any(idx >= scope_count - 1 for idx in positions)
            does_not_touch_direct_line = min(positions) >= direct
            if same_last_round2_total and touches_round2_last_line and does_not_touch_direct_line:
                continue

        # ขึ้นเมื่อคนในกลุ่มตีจบครบแล้ว ไม่ขึ้นก่อนกรรมการกดจบ
        if not all(athlete_round_status(row["athlete"], 1) == "finished" for row in group_rows):
            continue

        counts = [row.get("tiebreak_count", 0) for row in group_rows]
        totals = [row.get("tiebreak_total", 0) for row in group_rows]

        # ยังไม่ได้ยิงรอบพิเศษครบ หรือจำนวนเที่ยวไม่เท่ากัน = ต้องยิง/ยิงต่อ
        if min(counts) == 0 or len(set(counts)) > 1:
            result.update(row["athlete"].id for row in group_rows)
            continue

        # ยิงครบเท่ากันแล้วแต่คะแนนพิเศษยังเท่ากัน = ต้องยิงต่อ
        if len(set(totals)) == 1:
            result.update(row["athlete"].id for row in group_rows)

    return result


def overview_shootoff_ids(event: Event, round_no: int) -> set[int]:
    """ตรวจ Shoot-off สำหรับ Overview

    รอบ 1:
    - ใช้ TOTAL -> 5 -> 3 เพื่อจัดลำดับทุกคน
    - ถ้ายังเท่ากันในกลุ่มที่มีผลต่อ "เข้ารอบตรง" ต้อง Shoot-off
    - ข้อยกเว้นมีแค่เส้นสุดท้ายของสิทธิ์ตีรอบ 2: ถ้าคะแนนรวมเท่ากัน ให้ได้ตีรอบ 2 ทั้งหมด

    รอบ 2:
    - ต้องไม่มีลำดับซ้ำ เพราะต้องเอาไปเป็น Seed ต่อ
    - ใช้ SUM(R1+R2) -> จำนวน 5 รวม -> จำนวน 3 รวม
    - ถ้ายังเท่ากันหลังกลุ่มนั้นกดจบการตีครบ ให้ขึ้น Shoot-off ทันที
    """
    if round_no == 1:
        rows = build_round_ranking(event, 1)
        return round1_overview_unresolved_shootoff_ids(event, rows)

    if round_no == 2 and event.has_round_two:
        rows = [
            r for r in build_round_two_overview_rows(event)
            if not r.get("is_round2_direct_placeholder")
            and r.get("round2_has_played")
            and athlete_round_status(r["athlete"], 2) == "finished"
        ]
        rows.sort(key=lambda row: (
            -row.get("combined_total", row.get("total", 0)),
            -row.get("count_5", 0),
            -row.get("count_3", 0),
            -row.get("tiebreak_total", 0),
            row.get("display_order") if row.get("display_order") is not None else 999999,
            row["athlete"].id,
        ))
        # รอบ 2 จัดอันดับ/ตรวจ Shoot-off เฉพาะคนที่ตีจบแล้ว ไม่ลากคนที่ยังรอคิวมาคิด
        return overview_unresolved_shootoff_ids(rows, None, round_no=2)

    return set()

def build_round_two_overview_rows(event: Event) -> List[dict]:
    """หน้า Overview รอบ 2 แบบ Official + realtime

    แสดงรายชื่อทันทีตั้งแต่เปิดรอบ 2 โดยไม่ต้องรอคีย์ครบ
    รูปแบบข้อมูลในแต่ละแถวเก็บทั้งรอบ 1 และรอบ 2 เพื่อให้ตารางเดียวแสดง:
    Qualification Round 1 / Qualification Round 2 / SUM / RANKING
    """
    round1_rows = build_round_ranking(event, 1)
    direct_limit = direct_quota(event)
    lane_count = max(event.lane_count or 1, 1)

    direct_ids = exact_cut_ids(round1_rows, direct_quota(event))
    direct_pending_ids = unresolved_tie_ids(round1_rows, direct_quota(event))
    direct_rows = []
    for row in round1_rows:
        if row["athlete"].id in direct_ids:
            direct_row = {**row}
            direct_row["round_no"] = 2
            direct_row["round1_total"] = row["total"]
            direct_row["round2_total"] = None
            direct_row["combined_total"] = row["total"]
            direct_row["round1_by_station"] = row.get("by_station", {})
            direct_row["round2_by_station"] = None
            direct_row["is_round2_direct_placeholder"] = True
            direct_row["display_order"] = "-"
            direct_row["display_lane_no"] = "-"
            direct_row["display_lane_order"] = "-"
            direct_row["status"] = "direct"
            direct_rows.append(direct_row)

    # Overview รอบ 2 ต้องไม่มีลำดับซ้ำ: แสดงลำดับจริงตามตำแหน่งหลังตัดสินด้วย TOTAL -> 5 -> 3 -> Shoot-off
    for idx, direct_row in enumerate(direct_rows, start=1):
        direct_row["rank"] = idx
        direct_row["ordinal_rank"] = idx
        direct_row["view_order"] = idx
        direct_row["display_rank"] = idx

    # รายชื่อรอบ 2 ต้องมาจากสิทธิ์หลังรอบ 1 ไม่ใช่จากลายเซ็น/คะแนนรอบ 2
    round2_source_rows = [
        row for row in round1_rows
        if row["athlete"].id not in direct_ids
        and row["athlete"].id not in direct_pending_ids
        and is_round_two_candidate(event, row["athlete"])
    ]
    # ลำดับตีรอบ 2: คะแนนรอบ 1 น้อยสุดก่อน
    round2_source_rows.sort(key=lambda row: (
        row["total"], row["count_5"], row["count_3"], row["tiebreak_total"],
        row["athlete"].start_order if row["athlete"].start_order is not None else 999999,
        row["athlete"].id,
    ))

    round2_rows = []
    for idx, source_row in enumerate(round2_source_rows, start=1):
        athlete = source_row["athlete"]
        r2 = summarize_round(athlete.id, 2)
        combined_total = source_row["total"] + r2["total"]
        row = {
            "athlete": athlete,
            "round_no": 2,
            # หน้า Overview รอบ 2:
            # total = คะแนนรอบ 2 อย่างเดียว
            # combined_total = คะแนนรอบ 1 + รอบ 2
            "total": r2["total"],
            "combined_total": combined_total,
            "round1_total": source_row["total"],
            "round2_total": r2["total"],
            "count_5": source_row["count_5"] + r2["count_5"],
            "count_3": source_row["count_3"] + r2["count_3"],
            # รอบ 2 ต้องใช้ Shoot-off ของรอบ 2 เท่านั้นในการแยกอันดับ SUM
            # ห้ามเอา Shoot-off รอบ 1 มาบวก เพราะจะทำให้กดบันทึก Shoot-off รอบ 2 แล้วระบบยังมองว่าเที่ยวไม่เท่ากัน/ไม่ถูกบันทึก
            "tiebreak_total": r2["tiebreak_total"],
            "tiebreak_count": r2.get("tiebreak_count", 0),
            "round1_tiebreak_total": source_row.get("tiebreak_total", 0),
            "round1_tiebreak_count": source_row.get("tiebreak_count", 0),
            "round2_tiebreak_total": r2["tiebreak_total"],
            "round2_tiebreak_count": r2.get("tiebreak_count", 0),
            "status": athlete_round_status(athlete, 2),
            "by_station": r2["by_station"],
            "round1_by_station": source_row.get("by_station", {}),
            "round2_by_station": r2["by_station"],
            "red_cards": r2["red_cards"],
            "display_order": idx,
            "display_lane_no": ((idx - 1) % lane_count) + 1,
            "display_lane_order": ((idx - 1) // lane_count) + 1,
            "round1_rank": source_row["rank"],
            "is_round2_direct_placeholder": False,
        }
        round2_rows.append(row)

    # หน้า R2 ต้องเรียงคิวก่อน: คะแนนรอบ 1 ต่ำสุดในกลุ่มรอบ 2 ได้ตีก่อน
    # เมื่อมีคนเริ่มตี/มีคะแนนแล้ว จึงจัดอันดับเฉพาะคนที่ตีแล้วไว้ด้านบน
    played_rows = []
    waiting_rows = []
    for row in round2_rows:
        has_played = (
            row["status"] in {"active", "finished"}
            or row.get("round2_total", 0) > 0
            or row.get("tiebreak_count", 0) > 0
        )
        row["round2_has_played"] = has_played
        if has_played:
            played_rows.append(row)
        else:
            waiting_rows.append(row)

    played_rows.sort(key=lambda row: (
        -row["combined_total"],
        -row["count_5"],
        -row["count_3"],
        -row["tiebreak_total"],
        row["display_order"],
        row["athlete"].id,
    ))
    waiting_rows.sort(key=lambda row: (
        row["display_order"],
        row["athlete"].start_order if row["athlete"].start_order is not None else 999999,
        row["athlete"].id,
    ))

    base = len(direct_rows)
    for idx, row in enumerate(played_rows, start=1):
        real_rank = base + idx
        row["rank"] = real_rank
        row["ordinal_rank"] = real_rank
        row["display_rank"] = real_rank
        row["view_order"] = real_rank

    waiting_base = base + len(played_rows)
    for idx, row in enumerate(waiting_rows, start=1):
        real_rank = waiting_base + idx
        row["rank"] = real_rank
        row["ordinal_rank"] = real_rank
        row["display_rank"] = real_rank
        row["view_order"] = real_rank

    return direct_rows + played_rows + waiting_rows

def scorecard_print_positions() -> dict:
    return {
        "header": {
            "bib_no": {"left": 165, "top": 18},
            "name": {"left": 380, "top": 18},
            "affiliation": {"left": 925, "top": 18},
        },
        "rows": {
            1: {"top": 362},  # รอบที่ 1
            2: {"top": 437},  # รอบที่ 2
            3: {"top": 512},  # รอบ 16 คน (เมื่อเปิดใช้สาย 16 คน) หรือรอบ 8 คนในอีเวนต์เดิม
            4: {"top": 587},  # รอบ 8 คน หรือรอบรองชนะเลิศในอีเวนต์เดิม
            5: {"top": 662},  # รอบรองชนะเลิศ หรือรอบชิงชนะเลิศในอีเวนต์เดิม
            6: {"top": 737},  # รอบชิงชนะเลิศ เมื่อเปิดใช้สาย 16 คน
        },
        "station_cols": {
            1: {"6": 170, "7": 211, "8": 252, "9": 293, "total": 334},
            2: {"6": 378, "7": 419, "8": 460, "9": 501, "total": 542},
            3: {"6": 586, "7": 627, "8": 668, "9": 709, "total": 750},
            4: {"6": 794, "7": 835, "8": 876, "9": 917, "total": 958},
            5: {"6": 1002, "7": 1043, "8": 1084, "9": 1125, "total": 1166},
        },
        "right_cols": {
            "grand_total": 1262,
            "rank": 1350,
            "athlete_signature": 1448,
        },
        "signature_rows": {
            1: {"judge": 776, "recorder": 776},
            2: {"judge": 814, "recorder": 814},
            3: {"judge": 852, "recorder": 852},
            4: {"judge": 890, "recorder": 890},
            5: {"judge": 928, "recorder": 928},
            6: {"judge": 966, "recorder": 966},
        },
    }



def get_progression_groups(event: Event) -> dict:
    cache = _request_cache()
    cache_key = ("progression_groups", event.id)
    if has_request_context() and cache_key in cache:
        return cache[cache_key]
    round1_rows = build_round_ranking(event, 1)
    direct_ids = exact_cut_ids(round1_rows, direct_quota(event))
    round2_candidate_ids = set()
    passed_round2_ids = set()
    eliminated_ids = set()
    if event.has_round_two:
        round2_candidate_ids = round_two_candidate_ids(event) - direct_ids
        combined = build_combined_qualifiers(event)
        passed_round2_ids = {
            r["athlete"].id for r in combined
            if r.get("round2_total") is not None and r.get("passed_cut")
        }
    for athlete in event.athletes:
        aid = athlete.id
        if aid in direct_ids:
            continue
        if aid in round2_candidate_ids and aid not in passed_round2_ids:
            eliminated_ids.add(aid)
    result = {"direct": direct_ids, "round2_candidates": round2_candidate_ids, "round2_passed": passed_round2_ids, "eliminated": eliminated_ids}
    if has_request_context():
        cache[cache_key] = result
    return result


def compute_round_ranks(event: Event) -> dict[int, dict[int,int]]:
    result = {}
    for rn in [1,2]:
        result[rn] = {}
        for row in build_round_ranking(event, rn):
            result[rn][row["athlete"].id] = row["rank"]
    return result


def configured_bracket_start_round(event: Event) -> str:
    return {
        "รอบ 16 คน": "R16",
        "รอบ 8 คน": "QF",
        "รอบ 4 คน": "SF",
        "รอบรองชนะเลิศ": "SF",
    }.get(event.next_round_label, "QF")


def ensure_bracket(event: Event) -> list[BracketMatch]:
    """สร้างตาราง bracket ล่วงหน้าทุกรอบถึงรอบชิง

    - หน้าประกบคู่จะเห็นช่องรอครบตั้งแต่ต้น เช่น R16 -> QF -> SF -> Final
    - ยังไม่ดึงชื่อขึ้นรอบถัดไปจนกว่าจะบันทึกผู้ชนะ แต่กล่องรอจะแสดงไว้แล้ว
    - ไม่ขยายจำนวนคนเองจากกรณีคะแนนเท่ากันที่เส้นตัด
    """
    desired_round = configured_bracket_start_round(event)
    qualifiers = build_combined_qualifiers(event)
    seeds = [row for row in qualifiers if row.get("seed")]
    seed_total = len(seeds)
    if seed_total > 4 and desired_round == "SF":
        desired_round = "QF"

    existing = BracketMatch.query.filter_by(event_id=event.id).all()
    existing_rounds = {m.round_name for m in existing}
    initial_round = "R16" if "R16" in existing_rounds else ("QF" if "QF" in existing_rounds else ("SF" if "SF" in existing_rounds else None))
    if existing and initial_round != desired_round and not any(m.winner_id for m in existing):
        BracketMatch.query.filter_by(event_id=event.id).delete(synchronize_session=False)
        db.session.commit()
        existing = []

    orders = {
        "R16": [(1,16),(8,9),(5,12),(4,13),(3,14),(6,11),(7,10),(2,15)],
        "QF": [(1,8),(4,5),(3,6),(2,7)],
        "SF": [(1,4),(2,3)],
        "F": [(1,2)],
    }
    rounds_after = {"R16": ["R16", "QF", "SF", "F"], "QF": ["QF", "SF", "F"], "SF": ["SF", "F"]}
    existing_keys = {(m.round_name, m.match_no): m for m in BracketMatch.query.filter_by(event_id=event.id).all()}
    changed = False
    for round_name in rounds_after.get(desired_round, [desired_round, "F"]):
        for idx, pair in enumerate(orders[round_name], start=1):
            if (round_name, idx) in existing_keys:
                continue
            a_id = b_id = None
            if round_name == desired_round:
                a, b = pair
                arow = seeds[a - 1] if len(seeds) >= a else None
                brow = seeds[b - 1] if len(seeds) >= b else None
                a_id = arow["athlete"].id if arow else None
                b_id = brow["athlete"].id if brow else None
            db.session.add(BracketMatch(event_id=event.id, round_name=round_name, match_no=idx, athlete_a_id=a_id, athlete_b_id=b_id))
            changed = True
    if changed:
        db.session.commit()
    maybe_advance_bracket(event)
    return BracketMatch.query.filter_by(event_id=event.id).order_by(BracketMatch.round_name, BracketMatch.match_no).all()


def maybe_advance_bracket(event: Event) -> None:
    matches = BracketMatch.query.filter_by(event_id=event.id).all()
    r16 = sorted([m for m in matches if m.round_name == "R16"], key=lambda m: m.match_no)
    qf = sorted([m for m in matches if m.round_name == "QF"], key=lambda m: m.match_no)
    sf = sorted([m for m in matches if m.round_name == "SF"], key=lambda m: m.match_no)
    fn = sorted([m for m in matches if m.round_name == "F"], key=lambda m: m.match_no)
    changed = False

    def fill_match(match, a_id, b_id):
        nonlocal changed
        if match and match.winner_id is None:
            if match.athlete_a_id != a_id or match.athlete_b_id != b_id:
                match.athlete_a_id = a_id
                match.athlete_b_id = b_id
                changed = True

    if r16 and len(r16) >= 8 and qf and all(m.winner_id for m in r16):
        pairings = [(r16[0].winner_id, r16[1].winner_id), (r16[2].winner_id, r16[3].winner_id), (r16[4].winner_id, r16[5].winner_id), (r16[6].winner_id, r16[7].winner_id)]
        for idx, (a, b) in enumerate(pairings):
            fill_match(qf[idx], a, b)
    if qf and len(qf) >= 4 and sf and all(m.winner_id for m in qf):
        pairings = [(qf[0].winner_id, qf[1].winner_id), (qf[2].winner_id, qf[3].winner_id)]
        for idx, (a, b) in enumerate(pairings):
            fill_match(sf[idx], a, b)
    if sf and len(sf) >= 2 and fn and all(m.winner_id for m in sf):
        fill_match(fn[0], sf[0].winner_id, sf[1].winner_id)
    if changed:
        db.session.commit()


def sync_match_winner_from_scores(match: BracketMatch) -> None:
    event = Event.query.get(match.event_id)
    round_no = bracket_round_to_scorecard_round(match.round_name, event)
    if not match.athlete_a_id or not match.athlete_b_id:
        return
    sig_a = get_round_signature(match.athlete_a_id, round_no)
    sig_b = get_round_signature(match.athlete_b_id, round_no)
    if not (sig_a and sig_a.finished_at and sig_b and sig_b.finished_at):
        return
    total_a = summarize_round(match.athlete_a_id, round_no)["total"]
    total_b = summarize_round(match.athlete_b_id, round_no)["total"]
    if total_a == total_b:
        return
    match.winner_id = match.athlete_a_id if total_a > total_b else match.athlete_b_id
    db.session.commit()


def build_bracket_match_row(event: Event, athlete: Athlete | None, round_name: str, seed_map: dict[int, int]) -> dict:
    if not athlete:
        return {
            "athlete": None,
            "athlete_id": None,
            "team": "-",
            "name": "-",
            "r1": "-",
            "r1r2": "",
            "show_ref_score": False,
            "is_direct_qualifier": False,
            "stations": ["-"] * 5,
            "total": "-",
            "seed": "",
            "status": "waiting",
            "status_label": "รอคิว",
        }

    progression_groups = get_progression_groups(event)
    is_direct_qualifier = athlete.id in progression_groups.get("direct", set())

    r1 = summarize_round(athlete.id, 1)["total"]
    r2 = summarize_round(athlete.id, 2)["total"] if event.has_round_two else 0
    ref_total = r1 + (r2 or 0)
    show_ref_score = not is_direct_qualifier

    round_no = bracket_round_to_scorecard_round(round_name, event)
    current = summarize_round(athlete.id, round_no)
    round_status = athlete_round_status(athlete, round_no)

    return {
        "athlete": athlete,
        "athlete_id": athlete.id,
        "team": athlete.affiliation,
        "name": athlete.name,
        "r1": r1,
        "r1r2": ref_total if show_ref_score else "",
        "show_ref_score": show_ref_score,
        "is_direct_qualifier": is_direct_qualifier,
        "stations": [current["by_station"][station]["total"] for station in STATIONS],
        "total": current["total"],
        "seed": seed_map.get(athlete.id, ""),
        "round_no": round_no,
        "status": round_status,
        "status_label": {"waiting": "รอคิว", "active": "กำลังตี", "finished": "ตีเสร็จแล้ว"}.get(round_status, "รอคิว"),
    }

def generate_next_bib_no(event_id: int) -> str:
    count = Athlete.query.filter_by(event_id=event_id).count()
    return str(count + 1)


def recalculate_event_orders(event: Event) -> None:
    athletes = Athlete.query.filter_by(event_id=event.id).order_by(Athlete.start_order, Athlete.id).all()
    for idx, athlete in enumerate(athletes, start=1):
        athlete.start_order = idx
        athlete.bib_no = str(idx)
        athlete.lane_no = ((idx - 1) % event.lane_count) + 1
        athlete.lane_order = ((idx - 1) // event.lane_count) + 1


def reset_event_bracket(event: Event) -> None:
    BracketMatch.query.filter_by(event_id=event.id).delete()
    db.session.commit()


def normalize_header(value: str) -> str:
    return (value or '').strip().lower()


def parse_athletes_excel(file_storage) -> list[tuple[str, str]]:
    workbook = load_workbook(file_storage, data_only=True)
    sheet = workbook.active
    headers = [normalize_header(cell.value if cell.value is not None else '') for cell in sheet[1]]
    try:
        name_idx = headers.index('ชื่อ')
        affiliation_idx = headers.index('สังกัด')
    except ValueError as exc:
        raise ValueError('ไฟล์ Excel ต้องมีหัวคอลัมน์ชื่อ และ สังกัด') from exc

    rows: list[tuple[str, str]] = []
    for row_no, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
        name = str(row[name_idx]).strip() if len(row) > name_idx and row[name_idx] is not None else ''
        affiliation = str(row[affiliation_idx]).strip() if len(row) > affiliation_idx and row[affiliation_idx] is not None else ''
        if not name and not affiliation:
            continue
        if not name or not affiliation:
            raise ValueError(f'แถวที่ {row_no} ต้องมีทั้งชื่อและสังกัด')
        rows.append((name, affiliation))
    if not rows:
        raise ValueError('ไม่พบข้อมูลนักกีฬาในไฟล์ Excel')
    return rows

def dashboard_stats() -> dict:
    events_count = Event.query.count()
    athletes_count = Athlete.query.count()
    round1_rows = []
    for event in Event.query.all():
        round1_rows.extend(build_round_ranking(event, 1))
    top_score = max(round1_rows, key=lambda r: r["total"], default=None)
    affiliation_best = {}
    for row in round1_rows:
        key = row["athlete"].affiliation
        if key not in affiliation_best or row["total"] > affiliation_best[key]["total"]:
            affiliation_best[key] = row
    return {
        "events_count": events_count,
        "athletes_count": athletes_count,
        "top_score": top_score,
        "top_affiliations": sorted(affiliation_best.values(), key=lambda r: r["total"], reverse=True)[:8],
    }


@app.context_processor
def inject_globals():
    return {"now": datetime.now(), "MAX_RED_CARDS": MAX_RED_CARDS}


@app.route("/set-language/<lang>")
def set_language(lang):
    if lang in SUPPORTED_LANGS:
        session["lang"] = lang
    next_url = request.args.get("next") or request.referrer or url_for("index")
    return redirect(next_url)


@app.route("/")
def index():
    stats = dashboard_stats()
    events = Event.query.order_by(Event.competition_date.desc(), Event.id.desc()).all()
    return render_template("index.html", events=events, stats=stats)


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            login_user(user)
            flash("เข้าสู่ระบบสำเร็จ", "success")
            return redirect(url_for("index"))
        flash("ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง", "danger")
    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash("ออกจากระบบแล้ว", "info")
    return redirect(url_for("index"))


@app.route("/admin/users", methods=["GET", "POST"])
@login_required
@role_required("superadmin")
def manage_users():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        role = request.form.get("role", "user")
        if username and password and not User.query.filter_by(username=username).first():
            user = User(username=username, role=role)
            user.set_password(password)
            db.session.add(user)
            db.session.commit()
            flash("สร้างผู้ใช้สำเร็จ", "success")
        else:
            flash("สร้างผู้ใช้ไม่สำเร็จ กรุณาตรวจสอบข้อมูล", "danger")
    users = User.query.order_by(User.id).all()
    return render_template("users.html", users=users)


@app.route("/events/new", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def create_event():
    if request.method == "POST":
        event = Event(
            name=request.form["name"].strip(),
            event_group=request.form["event_group"],
            category=request.form["category"],
            competition_date=date.fromisoformat(request.form["competition_date"]),
            location=request.form.get("location", "").strip(),
            lane_count=int(request.form["lane_count"]),
            direct_qualifiers=int(request.form["direct_qualifiers"]),
            has_round_two=request.form.get("has_round_two") == "yes",
            round_two_cutoff_rank=int(request.form["round_two_cutoff_rank"]) if request.form.get("round_two_cutoff_rank") else None,
            next_round_label=request.form["next_round_label"],
            round_two_advancers=int(request.form.get("round_two_advancers") or 4),
            created_by=current_user.id,
        )
        db.session.add(event)
        db.session.commit()
        flash("สร้างอีเวนต์สำเร็จ", "success")
        return redirect(url_for("manage_athletes", event_id=event.id))
    return render_template("event_form.html")


@app.route("/events/<int:event_id>/edit", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def edit_event(event_id: int):
    event = Event.query.get_or_404(event_id)
    if request.method == "POST":
        event.name = request.form["name"].strip()
        event.event_group = request.form["event_group"]
        event.category = request.form["category"]
        event.competition_date = date.fromisoformat(request.form["competition_date"])
        event.location = request.form.get("location", "").strip()
        event.lane_count = int(request.form["lane_count"])
        event.direct_qualifiers = int(request.form["direct_qualifiers"])
        event.has_round_two = request.form.get("has_round_two") == "yes"
        event.round_two_cutoff_rank = int(request.form["round_two_cutoff_rank"]) if request.form.get("round_two_cutoff_rank") else None
        event.next_round_label = request.form["next_round_label"]
        event.round_two_advancers = int(request.form.get("round_two_advancers") or 4)
        recalculate_event_orders(event)
        db.session.commit()
        reset_event_bracket(event)
        flash("แก้ไขอีเวนต์สำเร็จ", "success")
        return redirect(url_for("event_overview", event_id=event.id, round=1))
    return render_template("event_form.html", event=event, is_edit=True)


@app.route("/events/<int:event_id>/delete", methods=["POST"])
@login_required
@role_required("admin", "superadmin")
def delete_event(event_id: int):
    event = Event.query.get_or_404(event_id)
    BracketMatch.query.filter_by(event_id=event.id).delete()
    db.session.delete(event)
    db.session.commit()
    flash("ลบอีเวนต์แล้ว", "info")
    return redirect(url_for("index"))


@app.route("/events/<int:event_id>/athletes", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def manage_athletes(event_id: int):
    event = Event.query.get_or_404(event_id)
    if request.method == "POST":
        name = request.form["name"].strip()
        affiliation = request.form["affiliation"].strip()
        next_order = Athlete.query.filter_by(event_id=event.id).count() + 1
        lane_no = ((next_order - 1) % event.lane_count) + 1
        lane_order = ((next_order - 1) // event.lane_count) + 1
        athlete = Athlete(
            event_id=event.id,
            bib_no=str(next_order),
            name=name,
            affiliation=affiliation,
            start_order=next_order,
            lane_no=lane_no,
            lane_order=lane_order,
            status="waiting",
        )
        db.session.add(athlete)
        db.session.commit()
        flash("เพิ่มนักกีฬาสำเร็จ", "success")
        return redirect(url_for("manage_athletes", event_id=event.id))
    athletes = Athlete.query.filter_by(event_id=event.id).order_by(Athlete.start_order).all()
    return render_template("athletes.html", event=event, athletes=athletes)


@app.route("/athletes/<int:athlete_id>/delete", methods=["POST"])
@login_required
@role_required("admin", "superadmin")
def delete_athlete(athlete_id: int):
    athlete = Athlete.query.get_or_404(athlete_id)
    event = athlete.event
    athlete_name = athlete.name
    db.session.delete(athlete)
    db.session.commit()
    recalculate_event_orders(event)
    sync_round_two_candidates(event)
    reset_event_bracket(event)
    db.session.commit()
    flash(f"ลบรายการ {athlete_name} เรียบร้อย", "success")
    return redirect(url_for("manage_athletes", event_id=event.id))


@app.route("/events/<int:event_id>/athletes/import", methods=["POST"])
@login_required
@role_required("admin", "superadmin")
def import_athletes_excel(event_id: int):
    event = Event.query.get_or_404(event_id)
    file = request.files.get("excel_file")
    if not file or not file.filename:
        flash("กรุณาเลือกไฟล์ Excel", "danger")
        return redirect(url_for("manage_athletes", event_id=event.id))
    if not file.filename.lower().endswith(".xlsx"):
        flash("รองรับเฉพาะไฟล์ .xlsx", "danger")
        return redirect(url_for("manage_athletes", event_id=event.id))
    try:
        rows = parse_athletes_excel(file)
        next_order = Athlete.query.filter_by(event_id=event.id).count() + 1
        for name, affiliation in rows:
            athlete = Athlete(
                event_id=event.id,
                bib_no=str(next_order),
                name=name,
                affiliation=affiliation,
                start_order=next_order,
                lane_no=((next_order - 1) % event.lane_count) + 1,
                lane_order=((next_order - 1) // event.lane_count) + 1,
                status="waiting",
            )
            db.session.add(athlete)
            next_order += 1
        db.session.commit()
        flash(f"นำเข้านักกีฬาสำเร็จ {len(rows)} คน", "success")
    except Exception as exc:
        db.session.rollback()
        flash(str(exc), "danger")
    return redirect(url_for("manage_athletes", event_id=event.id))


@app.route("/athletes-import-template.xlsx")
def athletes_import_template():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Athletes"
    sheet.append(["ชื่อ", "สังกัด"])
    sheet.append(["นายตัวอย่าง ใจดี", "ขอนแก่น"])
    sheet.append(["นางสาวตัวอย่าง แสนดี", "อุดรธานี"])
    stream = BytesIO()
    workbook.save(stream)
    stream.seek(0)
    return send_file(
        stream,
        as_attachment=True,
        download_name="athletes_import_template.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/events/<int:event_id>/athletes/randomize", methods=["POST"])
@login_required
@role_required("admin", "superadmin")
def randomize_athletes(event_id: int):
    import random

    event = Event.query.get_or_404(event_id)
    athletes = Athlete.query.filter_by(event_id=event.id).all()
    random.shuffle(athletes)
    for idx, athlete in enumerate(athletes, start=1):
        athlete.start_order = idx
        athlete.lane_no = ((idx - 1) % event.lane_count) + 1
        athlete.lane_order = ((idx - 1) // event.lane_count) + 1
    db.session.commit()
    flash("สุ่มลำดับใหม่แล้ว", "success")
    return redirect(url_for("manage_athletes", event_id=event.id))



@app.route("/events/<int:event_id>/overview")
def event_overview(event_id: int):
    event = Event.query.get_or_404(event_id)
    round_no = int(request.args.get("round", 1))
    if round_no == 2 and event.has_round_two:
        sync_round_two_candidates(event)
    if round_no == 2 and event.has_round_two:
        rows = build_round_two_overview_rows(event)
    else:
        rows = build_round_ranking(event, round_no)
    groups = get_progression_groups(event)
    shoot_off_ids = overview_shootoff_ids(event, round_no)
    for row in rows:
        aid = row["athlete"].id
        row["shoot_off_required"] = aid in shoot_off_ids
        row["shoot_off_group_ids"] = shootoff_group_ids(rows, aid, round_no) if row["shoot_off_required"] else [aid]
        if row["shoot_off_required"]:
            row["progress_class"] = "shoot-off-required"
        elif round_no == 1:
            # หน้า Overview รอบ 1 แยกแค่เข้ารอบตรง/ได้สิทธิ์ตีรอบ 2 ก่อน ไม่ลงสี Knockout ที่นี่
            row["progress_class"] = "qualified-direct" if aid in groups["direct"] else ("round2-candidate" if aid in groups["round2_candidates"] else "")
        else:
            row["progress_class"] = ("qualified-round2" if aid in groups["round2_passed"] else ("eliminated" if aid in groups["eliminated"] else ""))
        row["cut_line_after"] = False

    # เส้นแบ่งกลุ่มสำคัญบนหน้า Overview
    # รอบ 1: แยกคนเข้ารอบตรงออกจากคนมีสิทธิ์ตีรอบ 2
    # รอบ 2: ขีดใต้คนสุดท้ายที่ผ่านจากรอบ 2 เข้า Bracket
    if round_no == 1:
        last_direct_idx = None
        last_round2_candidate_idx = None
        for idx, row in enumerate(rows):
            aid = row["athlete"].id
            if aid in groups["direct"]:
                last_direct_idx = idx
            if aid in groups["round2_candidates"]:
                last_round2_candidate_idx = idx
        if last_direct_idx is not None and last_direct_idx < len(rows) - 1:
            rows[last_direct_idx]["cut_line_after"] = True
        if last_round2_candidate_idx is not None and last_round2_candidate_idx < len(rows) - 1:
            rows[last_round2_candidate_idx]["cut_line_after"] = True
    elif round_no == 2:
        last_passed_idx = None
        for idx, row in enumerate(rows):
            if row["athlete"].id in groups["round2_passed"]:
                last_passed_idx = idx
        if last_passed_idx is not None and last_passed_idx < len(rows) - 1:
            rows[last_passed_idx]["cut_line_after"] = True

    combined_rows = []
    return render_template(
        "overview.html",
        event=event,
        round_no=round_no,
        rows=rows,
        combined_rows=combined_rows,
        theme=event_theme(event.category),
        station_images=[f"station_{i}.png" for i in STATIONS],
        
    )


@app.route("/events/<int:event_id>/overview-data")
def overview_data(event_id: int):
    event = Event.query.get_or_404(event_id)
    round_no = int(request.args.get("round", 1))
    if round_no == 2 and event.has_round_two:
        sync_round_two_candidates(event)
    if round_no == 2 and event.has_round_two:
        rows = build_round_two_overview_rows(event)
    else:
        rows = build_round_ranking(event, round_no)
    groups = get_progression_groups(event)
    shoot_off_ids = overview_shootoff_ids(event, round_no)
    for row in rows:
        row["shoot_off_required"] = row["athlete"].id in shoot_off_ids
        row["shoot_off_group_ids"] = shootoff_group_ids(rows, row["athlete"].id, round_no) if row["shoot_off_required"] else [row["athlete"].id]
        row["cut_line_after"] = False
    if round_no == 1:
        last_direct_idx = None
        last_round2_candidate_idx = None
        for idx, row in enumerate(rows):
            aid = row["athlete"].id
            if aid in groups["direct"]:
                last_direct_idx = idx
            if aid in groups["round2_candidates"]:
                last_round2_candidate_idx = idx
        if last_direct_idx is not None and last_direct_idx < len(rows) - 1:
            rows[last_direct_idx]["cut_line_after"] = True
        if last_round2_candidate_idx is not None and last_round2_candidate_idx < len(rows) - 1:
            rows[last_round2_candidate_idx]["cut_line_after"] = True
    elif round_no == 2:
        last_passed_idx = None
        for idx, row in enumerate(rows):
            if row["athlete"].id in groups["round2_passed"]:
                last_passed_idx = idx
        if last_passed_idx is not None and last_passed_idx < len(rows) - 1:
            rows[last_passed_idx]["cut_line_after"] = True
    payload = []
    for row in rows:
        stations = {}
        for station in STATIONS:
            stations[str(station)] = {
                "6": row["by_station"][station]["distances"].get(6, 0),
                "7": row["by_station"][station]["distances"].get(7, 0),
                "8": row["by_station"][station]["distances"].get(8, 0),
                "9": row["by_station"][station]["distances"].get(9, 0),
                "total": row["by_station"][station]["total"],
            }
        round1_stations = None
        round2_stations = None
        if round_no == 2 and event.has_round_two:
            round1_stations = {str(st): (row.get("round1_by_station") or {}).get(st, {}).get("total", 0) for st in STATIONS}
            if row.get("round2_by_station") is not None:
                round2_stations = {str(st): row["round2_by_station"].get(st, {}).get("total", 0) for st in STATIONS}
            else:
                round2_stations = {str(st): None for st in STATIONS}
        aid = row["athlete"].id
        payload.append({
            "rank": row["rank"],
            "display_rank": row.get("display_rank", row["rank"]),
            "name": row["athlete"].name,
            "affiliation": row["athlete"].affiliation,
            "start_order": row["athlete"].start_order,
            "lane_no": row["athlete"].lane_no,
            "lane_order": row["athlete"].lane_order,
            "display_order": row.get("display_order", row["athlete"].start_order),
            "display_lane_no": row.get("display_lane_no", row["athlete"].lane_no),
            "display_lane_order": row.get("display_lane_order", row["athlete"].lane_order),
            "status": row["status"],
            "total": row.get("round2_total", row["total"]) if round_no == 2 and event.has_round_two and not row.get("is_round2_direct_placeholder") else row["total"],
            "round1_total": row.get("round1_total"),
            "round2_total": row.get("round2_total"),
            "combined_total": row.get("combined_total", row["total"]),
            "round1_stations": round1_stations,
            "round2_stations": round2_stations,
            "athlete_id": row["athlete"].id,
            "progress_class": ("shoot-off-required" if row.get("shoot_off_required") else (("qualified-direct" if aid in groups["direct"] else ("round2-candidate" if aid in groups["round2_candidates"] else "")) if round_no == 1 else ("qualified-round2" if aid in groups["round2_passed"] else ("eliminated" if aid in groups["eliminated"] else "")))),
            "shoot_off_required": row.get("shoot_off_required", False),
            "shoot_off_group_ids": row.get("shoot_off_group_ids", [row["athlete"].id]),
            "is_round2_direct_placeholder": row.get("is_round2_direct_placeholder", False),
            "round2_has_played": row.get("round2_has_played", False),
            "cut_line_after": row.get("cut_line_after", False),
            # ใช้สำหรับเรียงแถว realtime: รอบ 1 ต้องเรียงตามคะแนน/Rank, รอบ 2 ใช้ view_order ที่ build_round_two_overview_rows กำหนด
            "view_order": row.get("view_order", row["rank"]),
            "stations": stations,
        })
    return jsonify(payload)


@app.route("/events/<int:event_id>/overview-stats")
def overview_stats(event_id: int):
    """สถิติ 5/3 สำหรับเปิดดูประกอบการจัดลำดับ โดยไม่ทำให้ตาราง Overview หลักรก"""
    event = Event.query.get_or_404(event_id)
    round_no = int(request.args.get("round", 1))

    if round_no == 2 and event.has_round_two:
        rows = [r for r in build_round_two_overview_rows(event) if not r.get("is_round2_direct_placeholder")]
    else:
        rows = build_round_ranking(event, 1)

    # เรียงตามกติกาจริงที่ใช้ประกอบอันดับ: TOTAL -> 5 -> 3 -> Shoot-off
    rows = sorted(rows, key=lambda r: (
        -r.get("combined_total", r.get("total", 0)),
        -r.get("count_5", 0),
        -r.get("count_3", 0),
        -r.get("tiebreak_total", 0),
        r.get("display_order") if r.get("display_order") is not None else 999999,
        r["athlete"].id,
    ))

    data = []
    for idx, row in enumerate(rows, start=1):
        data.append({
            "rank": idx,
            "name": row["athlete"].name,
            "affiliation": row["athlete"].affiliation or "-",
            "total": row.get("combined_total", row.get("total", 0)),
            "count_5": row.get("count_5", 0),
            "count_3": row.get("count_3", 0),
            "status": ("ต้องตี Shoot-off" if row.get("shoot_off_required") else ""),
        })
    return jsonify(data)


@app.route("/api/scorecard/<int:athlete_id>/autosave", methods=["POST"])
@login_required
@role_required("admin", "superadmin")
def autosave_scorecard(athlete_id: int):
    athlete = Athlete.query.get_or_404(athlete_id)
    payload = request.get_json() or {}
    round_no = int(payload.get("round_no", 1))
    station_no = int(payload.get("station_no", 1))
    distance_m = int(payload.get("distance_m", 6))
    score_value = str(payload.get("score", "")).strip()
    red = bool(payload.get("red", False))
    ensure_round_entries(athlete.id, round_no)
    signature = ensure_signature(athlete.id, round_no)
    if not signature.started_at:
        signature.started_at = datetime.utcnow()
    athlete.status = "active"
    entry = ScoreEntry.query.filter_by(athlete_id=athlete.id, round_no=round_no, station_no=station_no, distance_m=distance_m).first()
    if entry:
        value = 0 if score_value == "" else int(score_value)
        value = max(0, min(5, value))
        entry.is_red_card = red
        entry.score = 0 if red else value
    db.session.commit()
    clear_request_cache()
    summary = summarize_round(athlete.id, round_no)
    station_entries = ScoreEntry.query.filter_by(
        athlete_id=athlete.id,
        round_no=round_no,
        station_no=station_no,
    ).all()
    station_red = sum(1 for e in station_entries if e.is_red_card)
    return jsonify({
        "ok": True,
        "station_total": summary["by_station"][station_no]["total"],
        "station_red": station_red,
        "round_total": summary["total"]
    })


@app.route("/athletes/<int:athlete_id>/scorecard", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def scorecard(athlete_id: int):
    athlete = Athlete.query.get_or_404(athlete_id)
    event = athlete.event
    round_no = int(request.args.get("round", 1))
    if round_no == 2 and not is_round_two_candidate(event, athlete):
        flash("นักกีฬาคนนี้ไม่มีสิทธิ์ตีรอบ 2", "warning")
        return redirect(url_for("event_overview", event_id=event.id, round=1))

        # GET: เปิดหน้า scorecard อย่างเดียว
    # ไม่สร้าง score_entry / ไม่เปลี่ยนสถานะ / ไม่ commit
    # POST หรือ autosave ค่อยสร้างข้อมูลจริง
    if request.method == "POST":
        ensure_round_entries(athlete.id, round_no)
        signature = ensure_signature(athlete.id, round_no)
    else:
        signature = ScoreSignature.query.filter_by(
            athlete_id=athlete.id,
            round_no=round_no
        ).first()

        if not signature:
            signature = ScoreSignature(
                athlete_id=athlete.id,
                round_no=round_no
            )

    if request.method == "POST":
        referee_name_key = f"referee_name_{round_no}"
        recorder_name_key = f"recorder_name_{round_no}"
        athlete_name_key = f"athlete_name_{round_no}"
        referee_sig_key = f"ref_sig_{round_no}"
        recorder_sig_key = f"rec_sig_{round_no}"
        athlete_sig_key = f"ath_sig_{round_no}"

        if not signature.started_at:
            signature.started_at = datetime.utcnow()
            athlete.status = "active"

        signature.recorder_name = request.form.get(recorder_name_key, "").strip() or signature.recorder_name
        signature.referee_name = request.form.get(referee_name_key, "").strip() or signature.referee_name
        signature.athlete_name = request.form.get(athlete_name_key, "").strip() or signature.athlete_name

        signature.recorder_signature = request.form.get(recorder_sig_key, "").strip() or signature.recorder_signature
        signature.referee_signature = request.form.get(referee_sig_key, "").strip() or signature.referee_signature
        signature.athlete_signature = request.form.get(athlete_sig_key, "").strip() or signature.athlete_signature

        bypass_code = request.form.get("bypass_code", "").strip()
        bypass_ok = current_user.role == "superadmin" or (current_user.role == "admin" and bypass_code == "7929")
        signed_ok = all([
            bool(signature.recorder_name or signature.recorder_signature),
            bool(signature.referee_name or signature.referee_signature),
            bool(signature.athlete_name or signature.athlete_signature),
        ])

        if bypass_ok or signed_ok:
            signature.bypass_signed = bypass_ok
            signature.finished_at = datetime.utcnow()
            athlete.status = "finished"
            db.session.commit()

            if round_no == 1 and event.has_round_two:
                sync_round_two_candidates(event)
                reset_event_bracket(event)
            elif round_no == 2 and event.has_round_two:
                reset_event_bracket(event)
            elif round_no >= 3:
                round_map = (
                    {3: "R16", 4: "QF", 5: "SF", 6: "F"}
                    if event_has_round_of_16(event)
                    else {3: "QF", 4: "SF", 5: "F"}
                )

                if round_no not in round_map:
                    flash("รอบแข่งขันไม่ถูกต้อง", "warning")
                    return redirect(url_for("bracket", event_id=event.id))

                match = BracketMatch.query.filter_by(
                    event_id=event.id,
                    round_name=round_map[round_no]
                ).filter(
                    (BracketMatch.athlete_a_id == athlete.id) |
                    (BracketMatch.athlete_b_id == athlete.id)
                ).first()

                if match:
                    sync_match_winner_from_scores(match)
                    maybe_advance_bracket(event)

                flash("จบการตีเรียบร้อย", "success")
                return redirect(url_for("bracket", event_id=event.id))

            flash("จบการตีเรียบร้อย", "success")
            return redirect(url_for("event_overview", event_id=event.id, round=round_no))

        flash("ต้องลงชื่ออย่างใดอย่างหนึ่ง (พิมพ์ชื่อหรือเขียน) ให้ครบทั้ง 3 ฝ่าย หรือใช้สิทธิ์ข้าม", "danger")
        return redirect(url_for("scorecard", athlete_id=athlete.id, round=round_no))

    template_data = build_scorecard_template_data(athlete.id)
    ranks = compute_round_ranks(event)
    template_data["round_ranks"] = {
        1: ranks.get(1, {}).get(athlete.id, ""),
        2: ranks.get(2, {}).get(athlete.id, ""),
        3: "",
        4: "",
        5: "",
        6: "",
    }

    round_station_running_totals = {}
    for rn in scorecard_round_numbers(event):
        running = {}
        acc = 0
        for st in [1, 2, 3, 4, 5]:
            val = template_data["station_totals"].get((rn, st), 0)
            acc += val
            running[st] = acc
        round_station_running_totals[rn] = running

    round_signatures = {}
    for rn in scorecard_round_numbers(event):
        round_signatures[rn] = get_round_signature(athlete.id, rn)

    combined_rows = build_combined_qualifiers(event) if event.has_round_two else []
    current_combined = next(
        (
            r for r in combined_rows
            if r.get("athlete") and r["athlete"].id == athlete.id
        ),
        None
    )

    display_order = athlete.start_order
    display_lane_no = athlete.lane_no
    display_lane_order = athlete.lane_order

    current_round_rows = build_round_ranking(event, round_no)
    current_row = next(
        (row for row in current_round_rows if row["athlete"].id == athlete.id),
        None
    )

    if current_row:
        display_order = current_row.get("display_order", display_order)
        display_lane_no = current_row.get("display_lane_no", display_lane_no)
        display_lane_order = current_row.get("display_lane_order", display_lane_order)

    return render_template(
        "scorecard.html",
        athlete=athlete,
        event=event,
        round_no=round_no,
        round_labels=scorecard_round_labels(event),
        score_map=template_data["score_map"],
        station_totals=template_data["station_totals"],
        station_reds=template_data["station_reds"],
        round_totals=template_data["round_totals"],
        round_ranks=template_data["round_ranks"],
        signature=signature,
        round_signatures=round_signatures,
        round_station_running_totals=round_station_running_totals,
        is_superadmin=(current_user.role == "superadmin"),
        theme=event_theme(event.category),
        combined_rows=combined_rows,
        current_combined=current_combined,
        station_images=[f"station_{i}.png" for i in STATIONS],
        display_order=display_order,
        display_lane_no=display_lane_no,
        display_lane_order=display_lane_order,
    )


@app.route("/events/<int:event_id>/scorecards-print-select", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def scorecards_print_select(event_id: int):
    event = Event.query.get_or_404(event_id)
    round_no = int(request.values.get("round", 1))

    athletes = athletes_for_scorecard_round(event, round_no)

    if request.method == "POST":
        print_mode = request.form.get("print_mode", "selected")
        selected_ids = request.form.getlist("athlete_ids")

        if print_mode == "all":
            return redirect(url_for(
                "scorecards_print_bulk",
                event_id=event.id,
                round=round_no
            ))

        if not selected_ids:
            flash("กรุณาเลือกนักกีฬาอย่างน้อย 1 คน", "warning")
            return redirect(url_for(
                "scorecards_print_select",
                event_id=event.id,
                round=round_no
            ))

        ids_text = ",".join(selected_ids)
        return redirect(url_for(
            "scorecards_print_bulk",
            event_id=event.id,
            round=round_no,
            ids=ids_text
        ))

    return render_template(
        "scorecards_print_select.html",
        event=event,
        athletes=athletes,
        round_no=round_no,
        round_labels=scorecard_round_labels(event),
    )


@app.route("/events/<int:event_id>/scorecards-print-bulk")
@login_required
@role_required("admin", "superadmin")
def scorecards_print_bulk(event_id: int):
    event = Event.query.get_or_404(event_id)
    round_no = int(request.args.get("round", 1))

    ids_text = request.args.get("ids", "").strip()
    selected_ids = []
    if ids_text:
        for raw_id in ids_text.split(","):
            raw_id = raw_id.strip()
            if raw_id.isdigit():
                selected_ids.append(int(raw_id))

    athletes = athletes_for_scorecard_round(event, round_no, selected_ids if selected_ids else None)

    print_items = [
        build_scorecard_print_context(athlete, round_no)
        for athlete in athletes
    ]

    return render_template(
        "scorecards_print_bulk.html",
        event=event,
        round_no=round_no,
        round_labels=scorecard_round_labels(event),
        print_items=print_items,
        station_images=[f"station_{i}.png" for i in STATIONS],
    )

@app.route("/athletes/<int:athlete_id>/scorecard-print")
@login_required
@role_required("admin", "superadmin")
def scorecard_print(athlete_id: int):
    athlete = Athlete.query.get_or_404(athlete_id)
    event = athlete.event
    round_no = int(request.args.get("round", 1))

    if round_no == 2 and event.has_round_two and not is_round_two_candidate(event, athlete):
        flash("นักกีฬาคนนี้ไม่มีสิทธิ์ตีรอบ 2", "warning")
        return redirect(url_for("event_overview", event_id=event.id, round=1))

    template_data = build_scorecard_template_data(athlete.id)
    ranks = compute_round_ranks(event)

    round_ranks = {
        1: ranks.get(1, {}).get(athlete.id, ""),
        2: ranks.get(2, {}).get(athlete.id, ""),
        3: "",
        4: "",
        5: "",
        6: "",
    }

    round_station_running_totals = {}
    for rn in scorecard_round_numbers(event):
        running = {}
        acc = 0
        for st in STATIONS:
            val = template_data["station_totals"].get((rn, st), 0)
            acc += val
            running[st] = acc
        round_station_running_totals[rn] = running

    round_signatures = {}
    for rn in scorecard_round_numbers(event):
        round_signatures[rn] = get_round_signature(athlete.id, rn)

    combined_rows = build_combined_qualifiers(event) if event.has_round_two else []
    current_combined = next(
        (r for r in combined_rows if r.get("athlete") and r["athlete"].id == athlete.id),
        None
    )

    display_order = athlete.start_order
    display_lane_no = athlete.lane_no
    display_lane_order = athlete.lane_order
    current_round_rows = build_round_ranking(event, round_no)
    current_row = next((row for row in current_round_rows if row["athlete"].id == athlete.id), None)
    if current_row:
        display_order = current_row.get("display_order", display_order)
        display_lane_no = current_row.get("display_lane_no", display_lane_no)
        display_lane_order = current_row.get("display_lane_order", display_lane_order)

    positions = scorecard_print_positions()

    return render_template(
        "scorecard_print.html",
        athlete=athlete,
        event=event,
        round_no=round_no,
        round_labels=scorecard_round_labels(event),
        score_map=template_data["score_map"],
        station_totals=template_data["station_totals"],
        station_reds=template_data["station_reds"],
        round_totals=template_data["round_totals"],
        round_ranks=round_ranks,
        round_signatures=round_signatures,
        round_station_running_totals=round_station_running_totals,
        current_combined=current_combined,
        combined_rows=combined_rows,
        display_order=display_order,
        display_lane_no=display_lane_no,
        display_lane_order=display_lane_order,
        positions=positions,
        station_images=[f"station_{i}.png" for i in STATIONS],
    )

@app.route("/athletes/<int:athlete_id>/activate", methods=["POST"])
@login_required
@role_required("admin", "superadmin")
def activate_scorecard(athlete_id: int):
    athlete = Athlete.query.get_or_404(athlete_id)
    round_no = int(request.args.get("round", 1))
    signature = ensure_signature(athlete.id, round_no)
    if not signature.finished_at:
        if not signature.started_at:
            signature.started_at = datetime.utcnow()
        athlete.status = "active"
        db.session.commit()
    return jsonify({"ok": True, "status": athlete_round_status(athlete, round_no)})


@app.route("/athletes/<int:athlete_id>/tiebreak", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def tiebreak(athlete_id: int):
    athlete = Athlete.query.get_or_404(athlete_id)
    round_no = int(request.values.get("round", request.args.get("round", 1)) or 1)
    if request.method == "POST":
        # ไม่ลบของเดิม: หากยังเท่ากัน ให้กลับมาตีเที่ยวพิเศษเพิ่มได้เรื่อย ๆ
        for station_no in STATIONS:
            score = int(request.form.get(f"tb_{station_no}", 0) or 0)
            db.session.add(TieBreakEntry(athlete_id=athlete.id, round_no=round_no, station_no=station_no, score=score))
        db.session.commit()
        clear_request_cache()
        flash("บันทึกผลเที่ยวพิเศษแล้ว ถ้ายังเท่ากันให้บันทึกเที่ยวพิเศษเพิ่มอีกครั้ง", "success")
        return redirect(url_for("event_overview", event_id=athlete.event_id, round=round_no))
    entries = TieBreakEntry.query.filter_by(athlete_id=athlete.id, round_no=round_no).all()
    existing = {station: sum(e.score for e in entries if e.station_no == station) for station in STATIONS}
    return render_template("tiebreak.html", athlete=athlete, round_no=round_no, existing=existing, athletes=[athlete], event=athlete.event)


@app.route("/events/<int:event_id>/tiebreak", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def event_tiebreak(event_id: int):
    event = Event.query.get_or_404(event_id)
    round_no = int(request.values.get("round", request.args.get("round", 1)) or 1)
    raw_ids = request.values.get("ids", "")
    athlete_ids = []
    for part in raw_ids.replace(" ", "").split(','):
        if part.isdigit():
            athlete_ids.append(int(part))
    athletes = Athlete.query.filter(Athlete.event_id == event.id, Athlete.id.in_(athlete_ids)).order_by(Athlete.start_order, Athlete.id).all() if athlete_ids else []
    if not athletes:
        flash("กรุณาเลือกนักกีฬาที่ต้องตี Shoot-off อย่างน้อย 2 คน", "warning")
        return redirect(url_for("event_overview", event_id=event.id, round=round_no))
    if request.method == "POST":
        entries_data = []
        for athlete in athletes:
            for station_no in STATIONS:
                score = int(request.form.get(f"tb_{athlete.id}_{station_no}", 0) or 0)
                entries_data.append({
                    "athlete_id": athlete.id,
                    "round_no": round_no,
                    "station_no": station_no,
                    "score": score,
                })

        try:
            for item in entries_data:
                db.session.add(TieBreakEntry(**item))
            db.session.commit()
        except Exception as exc:
            # PostgreSQL บางฐานที่ย้ายมาจาก SQLite มี tie_break_entry.id เป็น NOT NULL
            # แต่ไม่มี default sequence ทำให้ INSERT แล้ว id เป็น null
            db.session.rollback()
            msg = str(exc).lower()
            if "tie_break_entry" in msg and ("null value in column" in msg or "not-null constraint" in msg):
                next_id = next_manual_id(TieBreakEntry)
                if next_id is None:
                    raise
                for item in entries_data:
                    item["id"] = next_id
                    next_id += 1
                    db.session.add(TieBreakEntry(**item))
                db.session.commit()
            else:
                raise

        clear_request_cache()
        flash("บันทึก Shoot-off พร้อมกันแล้ว ถ้ายังเท่ากันให้เลือกกลุ่มเดิมแล้วบันทึกเพิ่ม", "success")
        return redirect(url_for("event_overview", event_id=event.id, round=round_no))
    existing = {}
    for athlete in athletes:
        entries = TieBreakEntry.query.filter_by(athlete_id=athlete.id, round_no=round_no).all()
        existing[athlete.id] = {station: sum(e.score for e in entries if e.station_no == station) for station in STATIONS}
    return render_template("tiebreak.html", event=event, athletes=athletes, athlete=athletes[0], round_no=round_no, existing=existing, bulk_mode=True, ids=','.join(str(a.id) for a in athletes))


@app.route("/events/<int:event_id>/bracket")
def bracket(event_id: int):
    event = Event.query.get_or_404(event_id)
    qualifiers = build_combined_qualifiers(event)
    ensure_bracket(event)
    maybe_advance_bracket(event)
    matches = BracketMatch.query.filter_by(event_id=event.id).order_by(BracketMatch.round_name, BracketMatch.match_no).all()
    athlete_map = {a.id: a for a in event.athletes}
    seed_map = {row["athlete"].id: row.get("seed") for row in qualifiers if row.get("athlete") and row.get("seed")}
    grouped = {"R16": [], "QF": [], "SF": [], "F": []}
    for m in matches:
        a = athlete_map.get(m.athlete_a_id)
        b = athlete_map.get(m.athlete_b_id)
        grouped.setdefault(m.round_name, []).append({
            "match": m,
            "a": build_bracket_match_row(event, a, m.round_name, seed_map),
            "b": build_bracket_match_row(event, b, m.round_name, seed_map),
            "winner": athlete_map.get(m.winner_id),
            "round_no": bracket_round_to_scorecard_round(m.round_name, event),
            "status": bracket_match_status(event, m),
        })
    combined_rows = build_combined_qualifiers(event) if event.has_round_two else qualifiers
    start_round = configured_bracket_start_round(event)
    return render_template("bracket.html", event=event, grouped=grouped, combined_rows=combined_rows, start_round=start_round)


@app.route("/matches/<int:match_id>/winner", methods=["POST"])
@login_required
@role_required("admin", "superadmin")
def set_match_winner(match_id: int):
    match = BracketMatch.query.get_or_404(match_id)
    winner_id = int(request.form.get("winner_id"))
    if winner_id not in {match.athlete_a_id, match.athlete_b_id}:
        flash("ผู้ชนะไม่ถูกต้อง", "danger")
        return redirect(url_for("bracket", event_id=match.event_id))
    match.winner_id = winner_id
    db.session.commit()
    maybe_advance_bracket(Event.query.get(match.event_id))
    flash("บันทึกผู้ชนะแล้ว", "success")
    return redirect(url_for("bracket", event_id=match.event_id))


@app.route("/events/<int:event_id>/bracket.xlsx")
def bracket_excel(event_id: int):
    event = Event.query.get_or_404(event_id)
    qualifiers = build_combined_qualifiers(event)
    wb = Workbook()
    ws = wb.active
    ws.title = "Bracket"
    ws.append(["Seed","Name","Affiliation","Round1","Round2","Sum"])
    for idx, row in enumerate(qualifiers, start=1):
        ws.append([idx, row["athlete"].name, row["athlete"].affiliation, row.get("round1_total", row["total"]), row.get("round2_total", ""), row["total"]])
    stream = BytesIO(); wb.save(stream); stream.seek(0)
    return send_file(stream, as_attachment=True, download_name=f"event_{event.id}_bracket.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.route("/events/<int:event_id>/bracket_data")
def bracket_data(event_id: int):
    event = Event.query.get_or_404(event_id)
    preload_event_score_data(event)

    maybe_advance_bracket(event)

    matches = BracketMatch.query.filter_by(event_id=event.id).order_by(
        BracketMatch.round_name,
        BracketMatch.match_no
    ).all()

    athlete_map = {a.id: a for a in event.athletes}
    qualifiers = build_combined_qualifiers(event)
    seed_map = {row["athlete"].id: row.get("seed", idx) for idx, row in enumerate(qualifiers, start=1) if row.get("athlete") and row.get("seed")}

    grouped = {"R16": [], "QF": [], "SF": [], "F": []}
    for m in matches:
        grouped.setdefault(m.round_name, []).append({
            "match_no": m.match_no,
            "round_no": bracket_round_to_scorecard_round(m.round_name, event),
            "winner_id": m.winner_id,
            "status": bracket_match_status(event, m),
            "a": build_bracket_row_data(event, athlete_map.get(m.athlete_a_id), m.round_name, seed_map),
            "b": build_bracket_row_data(event, athlete_map.get(m.athlete_b_id), m.round_name, seed_map),
        })

    return jsonify(grouped)


# =========================

# Results Approved report (SEA Games style)
# =========================
def _ra_text(value, default="") -> str:
    if value is None:
        return default
    return str(value)


def _ra_event_title(event: Event) -> str:
    raw = " ".join([event.event_group or "", event.category or "", event.name or ""]).strip()
    return (raw or event.name or "PÉTANQUE SHOOTING").upper()


def _ra_date_text(event: Event) -> str:
    try:
        return event.competition_date.strftime("%d %B %Y").upper()
    except Exception:
        return _ra_text(event.competition_date, "")


def _ra_default_setting(event: Event) -> SimpleNamespace:
    return SimpleNamespace(
        event_id=event.id,
        competition_title=(event.name or "33rd SEA GAMES THAILAND 2025").upper(),
        host_line=(event.location or "THAILAND : THE KINGDOM OF THAILAND").upper(),
        date_line=_ra_date_text(event),
        location_line=(event.location or "").upper(),
        country_label="COUNTRY",
        president_title="PRESIDENT OF WORLD PÉTANQUE AND BOULES FEDERATION",
        president_name="",
        technical_title="F.I.P.J.P VICE PRESIDENT OF\nPÉTANQUE TECHNICAL DELEGATE",
        technical_name="",
        umpires_text="",
        approved_text="……………………………APPROVED",
        show_official_pages=True,
        cover_main_logo_path=None,
        cover_bottom_logo_1_path=None,
        cover_bottom_logo_2_path=None,
        cover_bottom_logo_3_path=None,
        header_logo_1_path=None,
        header_logo_2_path=None,
        header_logo_3_path=None,
        header_logo_4_path=None,
        side_logo_path=None,
    )


def get_results_approved_setting(event: Event, create: bool = False):
    setting = ResultsApprovedSetting.query.filter_by(event_id=event.id).first()
    if setting:
        return setting
    default = _ra_default_setting(event)
    if not create:
        return default
    setting = ResultsApprovedSetting(
        event_id=event.id,
        competition_title=default.competition_title,
        host_line=default.host_line,
        date_line=default.date_line,
        location_line=default.location_line,
        country_label=default.country_label,
        president_title=default.president_title,
        president_name=default.president_name,
        technical_title=default.technical_title,
        technical_name=default.technical_name,
        umpires_text=default.umpires_text,
        approved_text=default.approved_text,
        show_official_pages=default.show_official_pages,
        cover_main_logo_path=default.cover_main_logo_path,
        cover_bottom_logo_1_path=default.cover_bottom_logo_1_path,
        cover_bottom_logo_2_path=default.cover_bottom_logo_2_path,
        cover_bottom_logo_3_path=default.cover_bottom_logo_3_path,
        header_logo_1_path=default.header_logo_1_path,
        header_logo_2_path=default.header_logo_2_path,
        header_logo_3_path=default.header_logo_3_path,
        header_logo_4_path=default.header_logo_4_path,
        side_logo_path=default.side_logo_path,
    )
    db.session.add(setting)
    db.session.commit()
    return setting


def _ra_setting_text(setting, attr: str, default: str = "") -> str:
    value = getattr(setting, attr, None)
    return _ra_text(value, default).strip() or default


RESULTS_APPROVED_LOGO_FIELDS = {
    "cover_main_logo": "cover_main_logo_path",
    "cover_bottom_logo_1": "cover_bottom_logo_1_path",
    "cover_bottom_logo_2": "cover_bottom_logo_2_path",
    "cover_bottom_logo_3": "cover_bottom_logo_3_path",
    "header_logo_1": "header_logo_1_path",
    "header_logo_2": "header_logo_2_path",
    "header_logo_3": "header_logo_3_path",
    "header_logo_4": "header_logo_4_path",
    "side_logo": "side_logo_path",
}

RESULTS_APPROVED_LOGO_DEFAULTS = {
    "cover_main": "results_approved_assets/thailand2025.png",
    "cover_bottom_1": "results_approved_assets/fipjp.png",
    "cover_bottom_2": "results_approved_assets/wpbf.png",
    "cover_bottom_3": "results_approved_assets/absc.png",
    "header_1": "results_approved_assets/absc.png",
    "header_2": "results_approved_assets/thailand2025.png",
    "header_3": "results_approved_assets/wpbf.png",
    "header_4": "results_approved_assets/fipjp.png",
    "side": "results_approved_assets/thailand2025.png",
}

ALLOWED_RESULTS_LOGO_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "gif"}


def _ra_logo_value(setting, attr: str, default: str) -> str:
    value = getattr(setting, attr, None)
    value = _ra_text(value).strip()
    return value or default


def _ra_logo_map(setting) -> dict:
    return {
        "cover_main": _ra_logo_value(setting, "cover_main_logo_path", RESULTS_APPROVED_LOGO_DEFAULTS["cover_main"]),
        "cover_bottom_1": _ra_logo_value(setting, "cover_bottom_logo_1_path", RESULTS_APPROVED_LOGO_DEFAULTS["cover_bottom_1"]),
        "cover_bottom_2": _ra_logo_value(setting, "cover_bottom_logo_2_path", RESULTS_APPROVED_LOGO_DEFAULTS["cover_bottom_2"]),
        "cover_bottom_3": _ra_logo_value(setting, "cover_bottom_logo_3_path", RESULTS_APPROVED_LOGO_DEFAULTS["cover_bottom_3"]),
        "header_1": _ra_logo_value(setting, "header_logo_1_path", RESULTS_APPROVED_LOGO_DEFAULTS["header_1"]),
        "header_2": _ra_logo_value(setting, "header_logo_2_path", RESULTS_APPROVED_LOGO_DEFAULTS["header_2"]),
        "header_3": _ra_logo_value(setting, "header_logo_3_path", RESULTS_APPROVED_LOGO_DEFAULTS["header_3"]),
        "header_4": _ra_logo_value(setting, "header_logo_4_path", RESULTS_APPROVED_LOGO_DEFAULTS["header_4"]),
        "side": _ra_logo_value(setting, "side_logo_path", RESULTS_APPROVED_LOGO_DEFAULTS["side"]),
    }


def _ra_save_uploaded_logo(event_id: int, field_name: str) -> str | None:
    uploaded = request.files.get(field_name)
    if not uploaded or not uploaded.filename:
        return None
    original = secure_filename(uploaded.filename)
    ext = original.rsplit(".", 1)[-1].lower() if "." in original else ""
    if ext not in ALLOWED_RESULTS_LOGO_EXTENSIONS:
        flash(f"ไฟล์โลโก้ {field_name} ต้องเป็น png, jpg, jpeg, webp หรือ gif", "danger")
        return None
    upload_dir = os.path.join(BASE_DIR, "static", "uploads", "results_approved", f"event_{event_id}")
    os.makedirs(upload_dir, exist_ok=True)
    filename = f"{field_name}_{datetime.utcnow().strftime('%Y%m%d%H%M%S%f')}.{ext}"
    uploaded.save(os.path.join(upload_dir, filename))
    return f"uploads/results_approved/event_{event_id}/{filename}"


def _ra_static_abs_path(static_filename: str | None) -> str | None:
    static_filename = _ra_text(static_filename).strip()
    if not static_filename:
        return None
    path = os.path.join(BASE_DIR, "static", *static_filename.split("/"))
    return path if os.path.exists(path) else None


def _ra_docx_add_center_image(doc, static_filename: str | None, width_inches: float = 1.35):
    path = _ra_static_abs_path(static_filename)
    if not path:
        return None
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_inches))
        return p
    except Exception:
        return None


def _ra_docx_add_logo_row(doc, static_filenames: list[str], width_inches: float = 0.75):
    paths = [_ra_static_abs_path(x) for x in static_filenames if x]
    paths = [p for p in paths if p]
    if not paths:
        return None
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, path in enumerate(paths):
            run = p.add_run()
            run.add_picture(path, width=Inches(width_inches))
            if idx != len(paths) - 1:
                p.add_run("     ")
        return p
    except Exception:
        return None


def _ra_split_name(full_name: str) -> tuple[str, str]:
    parts = [p for p in _ra_text(full_name).strip().split() if p]
    if len(parts) >= 2:
        return parts[0].upper(), " ".join(parts[1:]).upper()
    if parts:
        return "", parts[0].upper()
    return "", ""


def _ra_umpire_rows(setting) -> list[dict]:
    text = _ra_setting_text(setting, "umpires_text", "")
    rows = []
    for idx, line in enumerate([ln.strip() for ln in text.splitlines() if ln.strip()], start=1):
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 2:
            name, federation = parts[0], parts[1]
        else:
            name, federation = line, ""
        rows.append({"no": idx, "name": name.upper(), "federation": federation.upper()})
    return rows


def _ra_station_cells(summary: dict) -> list[int]:
    cells = []
    by_station = summary.get("by_station", {}) if summary else {}
    for station_no in STATIONS:
        cells.append(int(by_station.get(station_no, {}).get("total", 0) or 0))
    return cells


def _ra_distance_cells(summary: dict) -> list[int]:
    cells = []
    by_station = summary.get("by_station", {}) if summary else {}
    for station_no in STATIONS:
        distances = by_station.get(station_no, {}).get("distances", {})
        for distance_m in DISTANCES:
            cells.append(int(distances.get(distance_m, 0) or 0))
        cells.append(int(by_station.get(station_no, {}).get("total", 0) or 0))
    return cells


def _ra_station_groups(summary: dict) -> list[dict]:
    groups = []
    by_station = summary.get("by_station", {}) if summary else {}
    for station_no in STATIONS:
        station = by_station.get(station_no, {})
        distances = station.get("distances", {}) if station else {}
        values = [int(distances.get(distance_m, 0) or 0) for distance_m in DISTANCES]
        groups.append({
            "station": station_no,
            "values": values,
            "total": int(station.get("total", 0) or 0),
        })
    return groups


def _ra_athlete_rank(row: dict | None) -> str:
    if not row:
        return ""
    return _ra_text(row.get("display_rank") or row.get("ordinal_rank") or row.get("rank") or "")


def _ra_loser(match: BracketMatch):
    if not match or not match.winner_id:
        return None
    if match.athlete_a_id == match.winner_id:
        return Athlete.query.get(match.athlete_b_id) if match.athlete_b_id else None
    if match.athlete_b_id == match.winner_id:
        return Athlete.query.get(match.athlete_a_id) if match.athlete_a_id else None
    return None


def _ra_match_score(event: Event, match: BracketMatch, athlete_id: int | None) -> int | str:
    if not athlete_id:
        return ""
    round_no = bracket_round_to_scorecard_round(match.round_name, event)
    return summarize_round(athlete_id, round_no).get("total", 0)


def _ra_build_medal_rows(event: Event, bracket_matches: list[BracketMatch], fallback_rows: list[dict]) -> list[dict]:
    final = next((m for m in bracket_matches if m.round_name == "F"), None)
    semis = [m for m in bracket_matches if m.round_name == "SF"]
    medals = []

    if final and final.winner_id:
        gold = Athlete.query.get(final.winner_id)
        silver = _ra_loser(final)
        if gold:
            medals.append({"medal": "GOLD", "athlete": gold})
        if silver:
            medals.append({"medal": "SILVER", "athlete": silver})
        for semi in sorted(semis, key=lambda m: m.match_no):
            bronze = _ra_loser(semi)
            if bronze:
                medals.append({"medal": "BRONZE", "athlete": bronze})

    if medals:
        return medals[:4]

    labels = ["GOLD", "SILVER", "BRONZE", "BRONZE"]
    fallback = []
    for idx, row in enumerate(fallback_rows[:4]):
        athlete = row.get("athlete") if isinstance(row, dict) else None
        if athlete:
            fallback.append({"medal": labels[idx] if idx < len(labels) else "", "athlete": athlete})
    return fallback


def build_results_approved_context(event: Event) -> dict:
    """เตรียมข้อมูลรายงาน Results Approved SEA Games style จากข้อมูลจริงในระบบ"""
    preload_event_score_data(event)
    setting = get_results_approved_setting(event)
    country_label = _ra_setting_text(setting, "country_label", "COUNTRY").upper()
    approved_text = _ra_setting_text(setting, "approved_text", "……………………………APPROVED")
    athletes = sorted(event.athletes, key=lambda a: (a.start_order or 999999, a.id))

    round1_ranking = build_round_ranking(event, 1)
    round1_by_id = {row["athlete"].id: row for row in round1_ranking}

    entry_countries = []
    seen_country = set()
    for athlete in athletes:
        country = (athlete.affiliation or "").upper()
        key = country.strip().lower()
        if key and key not in seen_country:
            seen_country.add(key)
            entry_countries.append({"no": len(entry_countries) + 1, "country": country})

    name_rows = []
    qf1_rows = []
    qf1_detail_rows = []
    for athlete in athletes:
        family, given = _ra_split_name(athlete.name)
        name_rows.append({
            "no": athlete.start_order,
            "country": (athlete.affiliation or "").upper(),
            "family_name": family,
            "given_name": given,
            "name": (athlete.name or "").upper(),
        })
        row = round1_by_id.get(athlete.id)
        summary = summarize_round(athlete.id, 1)
        rank = _ra_athlete_rank(row)
        qf1_rows.append({
            "no": athlete.start_order,
            "country": (athlete.affiliation or "").upper(),
            "name": (athlete.name or "").upper(),
            "lane": athlete.lane_no,
            "points": summary.get("total", 0),
            "rank": rank,
        })
        qf1_detail_rows.append({
            "rank": rank,
            "country": (athlete.affiliation or "").upper(),
            "name": (athlete.name or "").upper(),
            "station_groups": _ra_station_groups(summary),
            "stations": _ra_station_cells(summary),
            "distance_cells": _ra_distance_cells(summary),
            "total": summary.get("total", 0),
        })

    qf2_rows = []
    qf2_detail_rows = []
    direct_rows = []
    if event.has_round_two:
        overview_r2 = build_round_two_overview_rows(event)
        direct_rows = [row for row in overview_r2 if row.get("is_round2_direct_placeholder")]
        round2_rows = [row for row in overview_r2 if not row.get("is_round2_direct_placeholder")]
        round2_rows = sorted(round2_rows, key=lambda row: (row.get("display_order") if isinstance(row.get("display_order"), int) else 999999, row["athlete"].id))
        for row in round2_rows:
            athlete = row["athlete"]
            r2_summary = summarize_round(athlete.id, 2)
            qf2_rows.append({
                "qf1_rank": row.get("round1_rank") or _ra_athlete_rank(round1_by_id.get(athlete.id)),
                "country": (athlete.affiliation or "").upper(),
                "name": (athlete.name or "").upper(),
                "lane": row.get("display_lane_no", ""),
                "r1": row.get("round1_total", 0),
                "r2": row.get("round2_total", 0),
                "total": row.get("combined_total", 0),
                "qf2_rank": _ra_athlete_rank(row),
            })
            qf2_detail_rows.append({
                "qf1_rank": row.get("round1_rank") or _ra_athlete_rank(round1_by_id.get(athlete.id)),
                "country": (athlete.affiliation or "").upper(),
                "name": (athlete.name or "").upper(),
                "station_groups": _ra_station_groups(r2_summary),
                "stations": _ra_station_cells(r2_summary),
                "distance_cells": _ra_distance_cells(r2_summary),
                "r1": row.get("round1_total", 0),
                "r2": row.get("round2_total", 0),
                "total": row.get("combined_total", 0),
                "qf2_rank": _ra_athlete_rank(row),
            })

    bracket_matches = BracketMatch.query.filter_by(event_id=event.id).order_by(BracketMatch.round_name, BracketMatch.match_no).all()
    bracket_round_order = {"R16": 1, "QF": 2, "SF": 3, "F": 4}
    bracket_matches = sorted(bracket_matches, key=lambda m: (bracket_round_order.get(m.round_name, 99), m.match_no))
    try:
        fallback_qualifiers = build_combined_qualifiers(event)
    except Exception:
        fallback_qualifiers = round1_ranking
    seed_map = {}
    for idx, row in enumerate(fallback_qualifiers or [], start=1):
        athlete = row.get("athlete") if isinstance(row, dict) else None
        if athlete:
            seed_map[athlete.id] = idx

    bracket_rows = []
    for match in bracket_matches:
        athlete_a = Athlete.query.get(match.athlete_a_id) if match.athlete_a_id else None
        athlete_b = Athlete.query.get(match.athlete_b_id) if match.athlete_b_id else None
        bracket_rows.append({
            "round_name": match.round_name,
            "round_label": {"R16": "ROUND OF 16", "QF": "QUARTERFINAL ROUND", "SF": "SEMIFINAL ROUND", "F": "FINAL ROUND"}.get(match.round_name, match.round_name),
            "match_no": match.match_no,
            "lane": match.match_no,
            "athlete_a": athlete_a,
            "athlete_b": athlete_b,
            "rank_a": seed_map.get(match.athlete_a_id, ""),
            "rank_b": seed_map.get(match.athlete_b_id, ""),
            "country_a": (athlete_a.affiliation if athlete_a else "").upper(),
            "country_b": (athlete_b.affiliation if athlete_b else "").upper(),
            "name_a": (athlete_a.name if athlete_a else "").upper(),
            "name_b": (athlete_b.name if athlete_b else "").upper(),
            "points_a": _ra_match_score(event, match, match.athlete_a_id),
            "points_b": _ra_match_score(event, match, match.athlete_b_id),
            "winner_id": match.winner_id,
        })

    medal_rows = _ra_build_medal_rows(event, bracket_matches, fallback_qualifiers or round1_ranking)
    medal_rows_out = []
    for r in medal_rows:
        athlete = r["athlete"]
        family, given = _ra_split_name(athlete.name)
        medal_rows_out.append({
            "medal": r["medal"],
            "athlete": athlete,
            "country": (athlete.affiliation or "").upper(),
            "family_name": family,
            "given_name": given,
        })

    return {
        "event": event,
        "setting": setting,
        "competition_title": _ra_setting_text(setting, "competition_title", event.name).upper(),
        "host_line": _ra_setting_text(setting, "host_line", event.location).upper(),
        "date_line": _ra_setting_text(setting, "date_line", _ra_date_text(event)).upper(),
        "location_line": _ra_setting_text(setting, "location_line", event.location).upper(),
        "country_label": country_label,
        "approved_text": approved_text,
        "event_title": _ra_event_title(event),
        "event_date_text": _ra_setting_text(setting, "date_line", _ra_date_text(event)).upper(),
        "athletes": athletes,
        "entry_countries": entry_countries,
        "name_rows": name_rows,
        "umpire_rows": _ra_umpire_rows(setting),
        "qf1_rows": qf1_rows,
        "qf1_detail_rows": qf1_detail_rows,
        "direct_rows": direct_rows,
        "qf2_rows": qf2_rows,
        "qf2_detail_rows": qf2_detail_rows,
        "bracket_rows": bracket_rows,
        "semifinal_rows": [r for r in bracket_rows if r["round_name"] == "SF"],
        "final_rows": [r for r in bracket_rows if r["round_name"] == "F"],
        "other_ko_rows": [r for r in bracket_rows if r["round_name"] not in {"SF", "F"}],
        "medal_rows": medal_rows_out,
        "logos": _ra_logo_map(setting),
        "stations": STATIONS,
        "distances": DISTANCES,
    }


def _ra_docx_set_cell_text(cell, text, bold=False, align="center", size_pt=9):
    cell.text = ""
    p = cell.paragraphs[0]
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align == "left" else WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_ra_text(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    try:
        from docx.shared import Pt
        run.font.size = Pt(size_pt)
    except Exception:
        pass


def _ra_docx_shade(cell, fill="D9D9D9"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _ra_docx_add_title(doc, text, size=18, spacing_after=4):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_ra_text(text))
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(spacing_after)
    return p


def _ra_docx_add_approved(doc, text="……………………………APPROVED"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p


def _ra_docx_add_table(doc, headers, rows, align_left_cols: set[int] | None = None, font_size=8):
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        _ra_docx_set_cell_text(table.rows[0].cells[i], header, bold=True, size_pt=font_size)
        _ra_docx_shade(table.rows[0].cells[i])
    align_left_cols = align_left_cols or set()
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            _ra_docx_set_cell_text(cells[i], value, align="left" if i in align_left_cols else "center", size_pt=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def make_results_approved_docx(event: Event) -> BytesIO:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ctx = build_results_approved_context(event)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    logos = ctx.get("logos", {})

    # Cover
    _ra_docx_add_center_image(doc, logos.get("cover_main"), width_inches=1.35)
    _ra_docx_add_title(doc, "SEA GAMES", 20)
    _ra_docx_add_title(doc, "THAILAND", 20)
    _ra_docx_add_title(doc, ctx["date_line"], 12)
    _ra_docx_add_title(doc, ctx["host_line"], 12)
    doc.add_paragraph()
    _ra_docx_add_title(doc, "PÉTANQUE", 26)
    _ra_docx_add_logo_row(doc, [logos.get("cover_bottom_1"), logos.get("cover_bottom_2"), logos.get("cover_bottom_3")], width_inches=0.75)
    doc.add_paragraph()
    _ra_docx_add_title(doc, "RESULTS", 24)
    _ra_docx_add_title(doc, "APPROVED", 10)
    doc.add_page_break()

    # Officials
    if getattr(ctx["setting"], "show_official_pages", True):
        _ra_docx_add_logo_row(doc, [logos.get("header_1"), logos.get("header_2"), logos.get("header_3"), logos.get("header_4")], width_inches=0.62)
        _ra_docx_add_title(doc, ctx["competition_title"], 14)
        _ra_docx_add_title(doc, _ra_setting_text(ctx["setting"], "president_title", "PRESIDENT"), 14)
        _ra_docx_add_title(doc, _ra_setting_text(ctx["setting"], "president_name", ""), 12)
        _ra_docx_add_title(doc, _ra_setting_text(ctx["setting"], "technical_title", "TECHNICAL DELEGATE"), 14)
        _ra_docx_add_title(doc, _ra_setting_text(ctx["setting"], "technical_name", ""), 12)
        _ra_docx_add_title(doc, "UMPIRE", 18)
        ump_rows = [[r["no"], r["name"], r["federation"]] for r in ctx["umpire_rows"]]
        if ump_rows:
            _ra_docx_add_table(doc, ["NO", "FAMILY NAME - GIVEN NAME", "FEDERATION"], ump_rows, align_left_cols={1}, font_size=9)
        _ra_docx_add_approved(doc, ctx["approved_text"])
        doc.add_page_break()

    # Event cover
    _ra_docx_add_title(doc, f"1. {ctx['event_title']}", 20)
    _ra_docx_add_title(doc, ctx["date_line"], 12)
    _ra_docx_add_table(doc, ["NO.", ctx["country_label"]], [[r["no"], r["country"]] for r in ctx["entry_countries"]], align_left_cols={1}, font_size=10)
    _ra_docx_add_approved(doc, ctx["approved_text"])
    doc.add_page_break()

    # Name list
    _ra_docx_add_title(doc, "NAME LISTS", 20)
    _ra_docx_add_title(doc, ctx["event_title"], 14)
    name_rows = [[r["no"], r["country"], r["family_name"], r["given_name"]] for r in ctx["name_rows"]]
    _ra_docx_add_table(doc, ["NO.", ctx["country_label"], "FAMILY NAME", "GIVEN NAME"], name_rows, align_left_cols={1,2,3}, font_size=9)
    _ra_docx_add_approved(doc, ctx["approved_text"])
    doc.add_page_break()

    # QF1 summary
    _ra_docx_add_title(doc, ctx["event_title"], 16)
    _ra_docx_add_title(doc, "QUALIFICATION ROUND 1", 16)
    qf1_rows = [[r["no"], r["country"], r["lane"], r["points"], r["rank"]] for r in ctx["qf1_rows"]]
    _ra_docx_add_table(doc, ["NO.", ctx["country_label"], "LANE", "POINTS", "RANK\n(QF1)"], qf1_rows, align_left_cols={1}, font_size=9)
    _ra_docx_add_approved(doc, ctx["approved_text"])
    doc.add_page_break()

    # QF1 detail landscape
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    _ra_docx_add_title(doc, f"{ctx['event_title']} - Qualification Shooting", 16)
    headers = ["Rank", ctx["country_label"], "Name"]
    for s in STATIONS:
        headers.extend([f"A{s} 6M", "7M", "8M", "9M", "Tot."])
    headers += ["Total", "Rank"]
    rows = []
    for r in ctx["qf1_detail_rows"]:
        rows.append([r["rank"], r["country"], r["name"], *r["distance_cells"], r["total"], r["rank"]])
    _ra_docx_add_table(doc, headers, rows, align_left_cols={1,2}, font_size=6)
    _ra_docx_add_approved(doc, ctx["approved_text"])
    doc.add_page_break()

    if ctx["qf2_rows"]:
        section = doc.add_section()
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        _ra_docx_add_title(doc, ctx["event_title"], 16)
        _ra_docx_add_title(doc, "QUALIFICATION ROUND 2", 16)
        qf2_rows = [[r["qf1_rank"], r["country"], r["lane"], r["r1"], r["r2"], r["total"], r["qf2_rank"]] for r in ctx["qf2_rows"]]
        _ra_docx_add_table(doc, ["RANK\n(QF1)", ctx["country_label"], "LANE", "R1", "R2", "TOTAL", "RANK\n(QF2)"], qf2_rows, align_left_cols={1}, font_size=9)
        _ra_docx_add_approved(doc, ctx["approved_text"])
        doc.add_page_break()

    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    _ra_docx_add_title(doc, "SEMIFINAL ROUND / FINAL ROUND", 18)
    ko_rows = []
    for r in ctx["semifinal_rows"] + ctx["final_rows"]:
        ko_rows.append([r["round_label"], r["match_no"], r["lane"], r["rank_a"], r["country_a"], r["points_a"]])
        ko_rows.append(["", "", "", r["rank_b"], r["country_b"], r["points_b"]])
    if ko_rows:
        _ra_docx_add_table(doc, ["ROUND", "MATCH", "LANE", "RANK", ctx["country_label"], "POINTS"], ko_rows, align_left_cols={4}, font_size=9)
    _ra_docx_add_approved(doc, ctx["approved_text"])
    doc.add_page_break()

    _ra_docx_add_title(doc, "RANKING RESULT", 20)
    _ra_docx_add_title(doc, ctx["event_title"], 16)
    medal_rows = [[r["medal"], r["country"], r["family_name"], r["given_name"]] for r in ctx["medal_rows"]]
    _ra_docx_add_table(doc, ["MEDAL", ctx["country_label"], "FAMILY NAME", "GIVEN NAME"], medal_rows, align_left_cols={1,2,3}, font_size=10)
    _ra_docx_add_approved(doc, ctx["approved_text"])

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


@app.route("/events/<int:event_id>/results-approved/settings", methods=["GET", "POST"])
@login_required
@role_required("admin", "superadmin")
def results_approved_settings(event_id: int):
    event = Event.query.get_or_404(event_id)
    setting = get_results_approved_setting(event, create=True)
    if request.method == "POST":
        setting.competition_title = request.form.get("competition_title", "").strip() or event.name
        setting.host_line = request.form.get("host_line", "").strip()
        setting.date_line = request.form.get("date_line", "").strip() or _ra_date_text(event)
        setting.location_line = request.form.get("location_line", "").strip()
        setting.country_label = request.form.get("country_label", "COUNTRY").strip().upper() or "COUNTRY"
        setting.president_title = request.form.get("president_title", "").strip()
        setting.president_name = request.form.get("president_name", "").strip()
        setting.technical_title = request.form.get("technical_title", "").strip()
        setting.technical_name = request.form.get("technical_name", "").strip()
        setting.umpires_text = request.form.get("umpires_text", "").strip()
        setting.approved_text = request.form.get("approved_text", "").strip() or "……………………………APPROVED"
        setting.show_official_pages = request.form.get("show_official_pages") == "yes"
        for field_name, attr in RESULTS_APPROVED_LOGO_FIELDS.items():
            if request.form.get(f"clear_{field_name}") == "yes":
                setattr(setting, attr, None)
            saved_logo = _ra_save_uploaded_logo(event.id, field_name)
            if saved_logo:
                setattr(setting, attr, saved_logo)
        db.session.commit()
        flash("บันทึกตั้งค่า Results Approved แล้ว", "success")
        return redirect(url_for("results_approved", event_id=event.id))
    logo_field_labels = [
        ("cover_main_logo", "โลโก้หลักบนปก"),
        ("cover_bottom_logo_1", "โลโก้ล่างปก 1"),
        ("cover_bottom_logo_2", "โลโก้ล่างปก 2"),
        ("cover_bottom_logo_3", "โลโก้ล่างปก 3"),
        ("header_logo_1", "โลโก้หัวกระดาษ 1"),
        ("header_logo_2", "โลโก้หัวกระดาษ 2"),
        ("header_logo_3", "โลโก้หัวกระดาษ 3"),
        ("header_logo_4", "โลโก้หัวกระดาษ 4"),
        ("side_logo", "โลโก้มุมซ้ายในหน้าผล"),
    ]
    return render_template(
        "results_approved_settings.html",
        event=event,
        setting=setting,
        logos=_ra_logo_map(setting),
        logo_field_labels=logo_field_labels,
    )


@app.route("/events/<int:event_id>/results-approved")
@login_required
def results_approved(event_id: int):
    event = Event.query.get_or_404(event_id)
    ctx = build_results_approved_context(event)
    return render_template("results_approved.html", **ctx)


@app.route("/events/<int:event_id>/results-approved.docx")
@login_required
def results_approved_docx(event_id: int):
    event = Event.query.get_or_404(event_id)
    try:
        stream = make_results_approved_docx(event)
    except ModuleNotFoundError:
        flash("ยังไม่ได้ติดตั้ง python-docx ให้รัน: python -m pip install -r requirements.txt", "warning")
        return redirect(url_for("results_approved", event_id=event.id))
    filename = f"results_approved_event_{event.id}.docx"
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

@app.route("/events/<int:event_id>/stats")
def event_stats(event_id: int):
    event = Event.query.get_or_404(event_id)
    round1_rows = build_round_ranking(event, 1)
    best = round1_rows[0] if round1_rows else None
    return render_template("stats.html", event=event, best=best, round1_rows=round1_rows)


# -----------------------------------------------------------------------------
# Public Live Report API / iframe support for Live Report Board
# -----------------------------------------------------------------------------
@app.after_request
def live_report_public_headers(response):
    response.headers.setdefault("Access-Control-Allow-Origin", "*")
    response.headers.setdefault("Access-Control-Allow-Headers", "Content-Type, Authorization")
    response.headers.setdefault("Access-Control-Allow-Methods", "GET, OPTIONS")
    # ให้ Report Board ฝังหน้า overview/bracket ผ่าน iframe ได้
    response.headers.pop("X-Frame-Options", None)
    response.headers.setdefault("Content-Security-Policy", "frame-ancestors *")
    return response


def _lr_shooting_event_payload(event):
    return {
        "id": event.id,
        "name": event.name,
        "event_group": event.event_group,
        "category": event.category,
        "competition_date": event.competition_date.isoformat() if event.competition_date else None,
        "location": event.location,
        "lane_count": event.lane_count,
        "athlete_count": len(event.athletes),
        "overview_url": url_for('event_overview', event_id=event.id, _external=True),
        "public_live_url": url_for('api_public_shooting_report', event_id=event.id, _external=True),
    }


def _lr_athlete_payload(event, athlete, rank=None):
    r1 = summarize_round(athlete.id, 1).get("total", 0)
    r2 = summarize_round(athlete.id, 2).get("total", 0)
    total = r1 + r2
    return {
        "rank": rank,
        "id": athlete.id,
        "bib_no": athlete.bib_no,
        "name": athlete.name,
        "affiliation": athlete.affiliation,
        "lane_no": athlete.lane_no,
        "lane_order": athlete.lane_order,
        "start_order": athlete.start_order,
        "status": athlete.status,
        "red_card_count": athlete.red_card_count,
        "round1_total": r1,
        "round2_total": r2,
        "total": total,
    }


@app.route('/api/public/shooting/events')
def api_public_shooting_events():
    events = Event.query.order_by(Event.created_at.desc(), Event.id.desc()).all()
    return jsonify({"ok": True, "source": "shooting", "events": [_lr_shooting_event_payload(e) for e in events]})


@app.route('/api/public/shooting/event/<int:event_id>/report')
def api_public_shooting_report(event_id: int):
    event = Event.query.get_or_404(event_id)
    preload_event_score_data(event)
    athletes = sorted(event.athletes, key=lambda a: (a.lane_no, a.lane_order, a.start_order))
    rows = [_lr_athlete_payload(event, a) for a in athletes]
    ranking = sorted(rows, key=lambda r: (r["total"], r["round1_total"], -r["red_card_count"]), reverse=True)
    for idx, row in enumerate(ranking, start=1):
        row["rank"] = idx
    return jsonify({
        "ok": True,
        "source": "shooting",
        "event": _lr_shooting_event_payload(event),
        "athletes": rows,
        "ranking": ranking,
        "live_url": url_for('event_overview', event_id=event.id, _external=True),
    })


@app.route('/public/shooting/<int:event_id>/live')
def public_shooting_live(event_id: int):
    # หน้า public สำหรับเอาไป iframe ใน Report Board โดยไม่ต้อง login
    event = Event.query.get_or_404(event_id)
    return redirect(url_for('event_overview', event_id=event.id, round=request.args.get('round', 1)))


def init_database_for_deploy() -> None:
    """Create database tables when running under gunicorn/Railway.

    Flask code inside __main__ is not executed by `gunicorn app:app`,
    so Railway needs this initialization during import.
    """
    os.makedirs(os.path.join(BASE_DIR, "instance"), exist_ok=True)
    with app.app_context():
        db.create_all()
        ensure_schema()
        seed_defaults()


# ให้ Railway/gunicorn สร้างตารางและ user ตั้งต้นทันทีตอน import app
init_database_for_deploy()


if __name__ == "__main__":
    app.run(
        host="0.0.0.0",
        port=int(os.environ.get("PORT", 8000)),
        debug=True
    )